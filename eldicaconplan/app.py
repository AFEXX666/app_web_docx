import os
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file, session
from jinja2 import Template
from flask_mysqldb import MySQL
from flask_wtf.csrf import CSRFProtect
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from config import config
from models.ModelUser import ModelUser, get_image_url
from models.entities.User import User
from datetime import datetime as t
from werkzeug.utils import secure_filename
from docx2pdf import convert
from docx import Document
import aspose.words as aw
from docx.shared import Pt
from collections import defaultdict
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import random
import string
from datetime import datetime
from PIL import Image
from io import BytesIO
from docx.shared import Inches
from docx.oxml.ns import qn
import base64
import docx
from docx.enum.text import  WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ALIGN_VERTICAL
import shutil
import webbrowser
import time
import openpyxl

app = Flask(__name__)

app.config['SECRET_KEY'] = 'B!1w8NAt1T^%kvhUI*S^'
csrf = CSRFProtect(app)
db = MySQL(app)
login_manager_app = LoginManager(app)


@login_manager_app.user_loader
def load_user(id):
    return ModelUser.get_by_id(db, id)


@app.route('/')
def index():
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        user = User(0, request.form['username'], request.form['password'])
        
        logged_user = ModelUser.login(db, user, table='users')

        if logged_user is None:
            logged_user = ModelUser.login(db, user, table='maestros')

        if logged_user is None:
            logged_user = ModelUser.login(db, user, table='superuser')
        
        if logged_user is not None:
            login_user(logged_user)
                    
            if logged_user.user_type == 'Administrador':
                return redirect(url_for('inicio'))
            elif logged_user.user_type == 'Maestro':
                return redirect(url_for('inicio'))
            elif logged_user.user_type == 'Superusuario':
                return redirect(url_for('home_sadmin'))
            else:
                flash("Tipo de usuario no válido...")
        else:
            flash("Credenciales incorrectas...")
        
        return render_template('auth/login.html')
    
    else:
        return render_template('auth/login.html')

@app.route('/logout')
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/home_user')
@login_required
def home_user():
    if current_user.user_type == "Maestro":
        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM `materias` ORDER BY idMateria DESC;")
        materia = cursor.fetchall()
        cursor = db.connection.cursor()
        cursor.execute("""
            SELECT plantillas.*, materias.nombre, users.Nombres, users.Apellidos
            FROM plantillas
            INNER JOIN materias ON plantillas.idMateria = materias.idMateria
            INNER JOIN users ON plantillas.idUser = users.idUser
            WHERE (idMaestro1 = %s OR idMaestro2 = %s OR idMaestro3 = %s)
            AND NOT EXISTS (
                SELECT 1
                FROM planeaciones
                WHERE planeaciones.idPlantilla = plantillas.idPlantilla
            )
            ORDER BY plantillas.idPlantilla DESC;
        """, (current_user.id, current_user.id, current_user.id))
        plantilla = cursor.fetchall()
        cursor = db.connection.cursor()
        cursor.execute("""
            SELECT planeaciones.*, users.Nombres, users.Apellidos, materias.nombre
            FROM planeaciones
            INNER JOIN plantillas ON planeaciones.idPlantilla = plantillas.idPlantilla
            INNER JOIN users ON plantillas.idUser = users.idUser
            INNER JOIN materias ON plantillas.idMateria = materias.idMateria
            WHERE (planeaciones.idMaestro = %s)
            ORDER BY idPlan DESC;
        """, (current_user.id,))
        plan = cursor.fetchall()
        cursor.close()
        print(plantilla)
        return render_template('user/home.html', materia=materia, plantilla=plantilla, plan=plan)
    else:
        return protected()

@app.route('/home_sadmin')
@login_required
def home_sadmin():
    if current_user.user_type == 'Superusuario':
        cursor = db.connection.cursor()
        cursor.execute("SELECT carrera.*, grado.nameGrado, edificio.nameE FROM carrera INNER JOIN grado ON carrera.idGrado = grado.idGrado INNER JOIN edificio ON carrera.idEdificio = edificio.idEdificio ORDER BY idCarrera DESC;")
        career = cursor.fetchall()

        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM `users` ORDER BY idUser DESC;")
        admin = cursor.fetchall()
        
        cursor = db.connection.cursor()
        cursor.execute("SELECT areas.*, carrera.carreraName FROM areas INNER JOIN carrera ON areas.idCarrera = carrera.idCarrera ORDER BY idArea DESC")
        special = cursor.fetchall()
        cursor.close()

        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM `edificio` ORDER BY idEdificio DESC;")
        edificio = cursor.fetchall()

        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM `grado` ORDER BY idGrado DESC;")
        grado = cursor.fetchall()

        

        cursor.close()
        return render_template('superAdmin/home.html', carrera=career, admin=admin, special=special, edificio=edificio, grado=grado)
    else:
        return protected()

@app.route('/home_admin')
@login_required
def home_admin():
    if current_user.user_type == "Administrador" and session['ass'] == 'plan':
        cursor = db.connection.cursor()
        cursor.execute("""SELECT plantillas.*, materias.nombre
        FROM plantillas
        INNER JOIN materias ON plantillas.idMateria = materias.idMateria
        JOIN maestros ms ON materias.idMaestro = ms.idMaestro
        JOIN grupos g ON ms.idGrupo = g.idGrupo
        JOIN grado gr ON g.idGrado = gr.idGrado
        JOIN areas a ON g.idArea = a.idArea
        JOIN carrera c ON a.idCarrera = c.idCarrera
        JOIN users u ON c.idCarrera = u.idCarrera
        WHERE plantillas.idUser = %s AND gr.nameGrado = %s AND a.idArea = %s
        ORDER BY plantillas.idPlantilla DESC;
        """,  (current_user.id, session.get('grado'), session.get('area'),))
        plantilla = cursor.fetchall()

        cursor = db.connection.cursor()
        cursor.execute("""
            SELECT a.idArea, a.aName
            FROM users AS u
            INNER JOIN carrera AS c ON u.idCarrera = c.idCarrera
            INNER JOIN areas AS a ON c.idCarrera = a.idCarrera
            INNER JOIN grado AS g ON c.idGrado = g.idGrado
            WHERE u.idUser = %s AND g.nameGrado = %s
        """, (current_user.id, session.get('grado'),))
        special = cursor.fetchall()

        cursor = db.connection.cursor()
        cursor.execute("""
            SELECT pl.*, ma.nombre , m.Nombres, m.ApellidosP
            FROM planeaciones AS pl
            JOIN plantillas AS p ON p.idPlantilla = pl.idPlantilla
            JOIN materias AS ma ON p.idMateria = ma.idMateria
            JOIN maestros AS m ON pl.idMaestro = m.idMaestro
            JOIN grupos g ON m.idGrupo = g.idGrupo
            JOIN grado gr ON g.idGrado = gr.idGrado
            JOIN areas a ON g.idArea = a.idArea
            JOIN carrera c ON a.idCarrera = c.idCarrera
            JOIN users u ON c.idCarrera = u.idCarrera
            WHERE p.idUser = %s AND gr.nameGrado = %s AND a.idArea = %s
            ORDER BY pl.idPlan DESC;
        """,  (current_user.id, session.get('grado'), session.get('area'),))
        plan = cursor.fetchall()

        cursor = db.connection.cursor()
        cursor.execute("""SELECT g.*, a.aName, gr.nameGrado AS nombreGrado
        FROM grupos AS g
        INNER JOIN areas AS a ON g.idArea = a.idArea
        INNER JOIN carrera AS c ON a.idCarrera = c.idCarrera
        INNER JOIN users AS u ON c.idCarrera = u.idCarrera
        INNER JOIN grado AS gr ON g.idGrado = gr.idGrado
        WHERE u.idUser = %s AND gr.nameGrado = %s AND a.idArea = %s
        ORDER BY g.idGrupo DESC;
        """, (current_user.id,  session.get('grado'),  session.get('area'),))

        grupo = cursor.fetchall()

        cursor = db.connection.cursor()
        cursor.execute("""SELECT m.*, a.aName AS nombre_area
            FROM materias m
            JOIN maestros ms ON m.idMaestro = ms.idMaestro 
            JOIN grupos g ON ms.idGrupo = g.idGrupo
            JOIN grado gr ON g.idGrado = gr.idGrado
            JOIN areas a ON g.idArea = a.idArea
            JOIN carrera c ON a.idCarrera = c.idCarrera
            JOIN users u ON c.idCarrera = u.idCarrera
            WHERE u.idUser = %s AND gr.nameGrado = %s AND a.idArea = %s
            ORDER BY m.idMateria DESC;
            """, (current_user.id, session.get('grado'), session.get('area'),))
        materia = cursor.fetchall()
        
        cursor = db.connection.cursor()
        cursor.execute("""SELECT ma.*, mt.nombre 
                       FROM maestros AS ma
                        JOIN materias mt ON ma.idMaestro = mt.idMaestro
                        JOIN grupos g ON ma.idGrupo = g.idGrupo
                        JOIN grado gr ON g.idGrado = gr.idGrado
                        JOIN areas a ON g.idArea = a.idArea
                        JOIN carrera c ON a.idCarrera = c.idCarrera
                        JOIN users u ON c.idCarrera = u.idCarrera
                       WHERE u.idUser = %s AND gr.nameGrado = %s AND a.idArea = %s 
                       ORDER BY idMaestro DESC; """, (current_user.id, session.get('grado'), session.get('area'),))
        maestros = cursor.fetchall()


        return render_template('admin/home.html',grupo=grupo,  materia=materia, maestros=maestros, plantilla=plantilla, special=special, plan=plan)
    else:
        return protected()
    
@app.route('/home_admin_dica')
@login_required
def home_admin_dica():
    if current_user.user_type == "Administrador" and session['ass'] == 'dica':
    
        cursor = db.connection.cursor()
        cursor.execute("""
            SELECT a.idArea, a.aName
            FROM users AS u
            INNER JOIN carrera AS c ON u.idCarrera = c.idCarrera
            INNER JOIN areas AS a ON c.idCarrera = a.idCarrera
            INNER JOIN grado AS g ON c.idGrado = g.idGrado
            WHERE u.idUser = %s AND g.nameGrado = %s
        """, (current_user.id, session.get('grado'),))
        special = cursor.fetchall()

        cursor.execute("""SELECT ma.* 
                       FROM students AS ma
                        JOIN grupos g ON ma.idGrupo = g.idGrupo
                        JOIN grado gr ON g.idGrado = gr.idGrado
                        JOIN areas a ON g.idArea = a.idArea
                        JOIN carrera c ON a.idCarrera = c.idCarrera
                        JOIN users u ON c.idCarrera = u.idCarrera
                       WHERE u.idUser = %s AND gr.nameGrado = %s AND a.idArea = %s 
                       ORDER BY idStudent DESC; """, (current_user.id, session.get('grado'), session.get('area'),))
        alumnos = cursor.fetchall()


        return render_template('dica/admin/home.html', special=special, alumnos=alumnos)
    else:
        return protected()
    

@app.route('/roll', methods=['GET'])
@login_required
def roll():
    session['grado'] = request.args.get('otro_dato')
    session['area'] = request.args.get('area')
    session['rol'] = request.args.get('rol')
    session['ass'] = request.args.get('ass')
    session['id'] = request.args.get('id')
    print(session['grado'])
    if session['rol'] == "Maestro" and session['ass'] == 'plan':
        return redirect(url_for('home_user'))
    elif session['rol'] == "Administrador" and session['ass'] == 'plan':
         return redirect(url_for('home_admin'))
    elif session['rol'] == "Administrador" and session['ass'] == 'dica':
         return redirect(url_for('home_admin_dica'))
    elif session['rol'] == "Maestro" and session['ass'] == 'dica':
         return redirect(url_for('home_user_dica'))
    elif session['rol'] == "dica":
         return redirect(url_for('home_user_dica'))
    else:
         return protected()
     
@app.route('/unidad', methods=['GET'])
@login_required
def unidad():
    session['unidad'] = request.args.get('uni')

    print(session['unidad'])
    return redirect(url_for('activity'))
     
@app.route('/gru', methods=['GET'])
@login_required
def gru():
    if current_user.user_type == 'Maestro':
        session['idMateria'] = request.args.get('mate')
        session['id'] = request.args.get('id')
        session['unidad'] = request.args.get('unidad')
        print(session['id'])
        return redirect(url_for('activity'))
    else:
        return protected()


@app.route('/home_user_dica')
@login_required
def home_user_dica():
    if current_user.user_type == "Maestro" and session['ass'] == 'dica':
        cursor = db.connection.cursor()
        cursor.execute("""SELECT m.*, a.aName AS nombre_area, g.nombre AS grade, g.cuatrimestre, g.idGrupo AS hinti
            FROM materias m
            JOIN maestros ms ON m.idMaestro = ms.idMaestro 
            JOIN grupos g ON ms.idGrupo = g.idGrupo
            JOIN grado gr ON g.idGrado = gr.idGrado
            JOIN areas a ON g.idArea = a.idArea
            JOIN carrera c ON a.idCarrera = c.idCarrera
            JOIN users u ON c.idCarrera = u.idCarrera
            WHERE m.idMaestro = %s AND gr.nameGrado = %s AND a.idArea = %s
            ORDER BY m.idMateria DESC;
            """, (current_user.id, session.get('grado'), session.get('area'),))
        materia = cursor.fetchall()
        return render_template('dica/subjects.html', materia=materia)
    else:
        return protected()
    

@app.route('/register')
@login_required
def register():
    if current_user.user_type == "Administrador" and session['ass'] == 'plan':
        cursor = db.connection.cursor()
        cursor.execute("""SELECT ma.* 
                       FROM maestros AS ma
                        JOIN grupos g ON ma.idGrupo = g.idGrupo
                        JOIN grado gr ON g.idGrado = gr.idGrado
                        JOIN areas a ON g.idArea = a.idArea
                        JOIN carrera c ON a.idCarrera = c.idCarrera
                        JOIN users u ON c.idCarrera = u.idCarrera
                       WHERE u.idUser = %s AND gr.nameGrado = %s AND a.idArea = %s 
                       ORDER BY idMaestro DESC; """, (current_user.id, session.get('grado'), session.get('area'),))
        maestros = cursor.fetchall()
        cursor = db.connection.cursor()
        cursor.execute("""SELECT g.*, a.aName, gr.nameGrado AS nombreGrado
        FROM grupos AS g
        INNER JOIN areas AS a ON g.idArea = a.idArea
        INNER JOIN carrera AS c ON a.idCarrera = c.idCarrera
        INNER JOIN users AS u ON c.idCarrera = u.idCarrera
        INNER JOIN grado AS gr ON g.idGrado = gr.idGrado
        WHERE u.idUser = %s AND gr.nameGrado = %s AND a.idArea = %s
        ORDER BY g.idGrupo DESC;
        """, (current_user.id,  session.get('grado'),  session.get('area'),))
        grupo = cursor.fetchall()
        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM materias",)
        materias = cursor.fetchall()
        print(maestros)
        return render_template('admin/register.html', maestros=maestros, grupo=grupo, materias=materias)
    
    else:
        return protected()
    
@app.route('/planeacion', methods=['POST'])
@login_required
def planeacion():
    if current_user.user_type == "Maestro":
        
        data = request.form
        
        name = data.get('name')

        def buscar_porcentajes_y_reemplazar(doc_path):
            timestamp = t.now().strftime('%Y%m%d%H%M%S') 
            filename_with_timestamp = timestamp + '_' + name
            carpeta_modificados = os.path.join('static/planeacion/', filename_with_timestamp)
            cuadros_encontrados = 0
            resultados = []
            porcentaje = [data.get(f'porcentaje_{i}', '') for i in range(1, 10)]
            actividades = [data.get(f'actividad_{i}', '') for i in range(1, 10)]
            semana = [data.get(f'numero_semana_{i}', '') for i in range(1, 15)]
            fecha_inicio = [data.get(f'fecha_inicio_{i}', '') for i in range(1, 15)]
            fecha_final = [data.get(f'fecha_fin_{i}', '') for i in range(1, 15)]
            conceptos = [data.get(f'conceptos_t_{i}', '') for i in range(1, 15)]
            temas_a_usar = [data.get(f'temas_{i}', '') for i in range(1, 15)]
            firma_base64 = data.get('firmas', '')  
            observacion = data.get('observacion', '')
            conocimientos = data.get('conocimientos', '')
            print(actividades)
            print(porcentaje)
            
            doc = docx.Document(doc_path)


            nueva_celda = None  
            
            for i, table in enumerate(doc.tables, 1):
                for row_index, row in enumerate(table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        if "% de la evaluación de la unidad" in cell.text.lower():
                            cuadros_encontrados = 0
                            for j in range(1, len(porcentaje) + 1):
                                if row_index + j < len(table.rows):
                                    siguiente_fila = table.rows[row_index + j]
                                    siguiente_celda = siguiente_fila.cells[cell_index]

                                    if cuadros_encontrados < len(porcentaje) and porcentaje[cuadros_encontrados]:
                                        siguiente_celda.text = porcentaje[cuadros_encontrados] + "\n"
                                    else:
                                        break

                                    for paragraph in siguiente_celda.paragraphs:
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        for run in paragraph.runs:
                                            run.font.name = 'Arial'

                                    cuadros_encontrados += 1
                                else:
                                    break

                            while cuadros_encontrados < len(porcentaje):
                                valor = porcentaje[cuadros_encontrados]
                                if valor:
                                    nueva_fila = table.add_row().cells
                                    nueva_celda = nueva_fila[cell_index]  
                                    nueva_celda.text = valor + "\n"
                                    for paragraph in nueva_celda.paragraphs:
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        for run in paragraph.runs:
                                            run.font.name = 'Arial'
                                cuadros_encontrados += 1

            for i, table in enumerate(doc.tables, 1):
                for row_index, row in enumerate(table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        if "descripción de la actividad para alcanzar el resultado de aprendizaje de la unidad" in cell.text.lower():
                            cuadros_encontrados = 0
                            for j in range(1, len(actividades) + 1):
                                if row_index + j < len(table.rows):
                                    siguiente_fila = table.rows[row_index + j]
                                    if siguiente_fila.cells[cell_index] == nueva_celda:  
                                        pass
                                    indice_texto = (j - 1) % len(actividades)
                                    texto_cuadro_debajo = siguiente_fila.cells[cell_index].text
                                    if actividades[indice_texto] != '':
                                        siguiente_fila.cells[cell_index].text = actividades[indice_texto]
                                        for paragraph in siguiente_fila.cells[cell_index].paragraphs:
                                            for run in paragraph.runs:
                                                run.font.name = 'Arial'
                                    else:
                                        siguiente_fila.cells[cell_index].text = ""
                                else:
                                    break


            cuadros_encontrados_totales = 0

            for i, table in enumerate(doc.tables, 1):
                for row_index, row in enumerate(table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        
                        if "fecha planeada por semana".lower() in cell.text.lower():
                            cuadros_encontrados_totales += 1
                            
                            texto_cuadros_debajo = []
                            for j in range(1, 5):  
                                if row_index + j < len(table.rows):
                                    siguiente_fila = table.rows[row_index + j]
                                    texto_cuadro_debajo = siguiente_fila.cells[cell_index].text
                                    texto_cuadros_debajo.append(texto_cuadro_debajo)
                                    siguiente_fila.cells[cell_index].text = ""  

                            resultados.append((cuadros_encontrados_totales, texto_cuadros_debajo))

                            
                            for j, texto_reemplazo in enumerate(semana):
                                if j < len(texto_cuadros_debajo):  
                                    siguiente_fila = table.rows[row_index + j + 1]
                                    siguiente_celda = siguiente_fila.cells[cell_index]
                                    if texto_reemplazo != '' and fecha_inicio[j] != '' and fecha_final[j] != '' and semana[j] != '':
                                        siguiente_celda.text = f"Semana {semana[j]}\n{fecha_inicio[j]}\nA\n{fecha_final[j]}"
                                        parrafo = siguiente_celda.paragraphs[0]
                                        run = parrafo.runs[0]
                                        font = run.font
                                        font.name = 'Arial'
                                        parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                        siguiente_celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                    else:
                                        siguiente_celda.text = ""

                            semana = semana[len(texto_cuadros_debajo):]  
                            fecha_inicio = fecha_inicio[len(texto_cuadros_debajo):]  
                            fecha_final = fecha_final[len(texto_cuadros_debajo):]  

            for i, table in enumerate(doc.tables, 1):
                for row_index, row in enumerate(table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        
                        if "observaciones a la materia" in cell.text.lower():
                            cuadros_encontrados += 1
                            
                            texto_cuadros_debajo = []
                            for j in range(1, 2):  
                                if row_index + j < len(table.rows):
                                    siguiente_fila = table.rows[row_index + j]
                                    siguiente_fila.cells[cell_index].text = observacion
                                    for paragraph in siguiente_fila.cells[cell_index].paragraphs:
                                        for run in paragraph.runs:
                                            run.font.name = 'Arial'
                                    
                            
                            resultados.append((i, texto_cuadros_debajo))

            def buscar_temas_en_tabla(doc_path):
                cuadros_encontrados = 0
                resultados = []

                doc = docx.Document(doc_path)

                for i, table in enumerate(doc.tables, 1):
                    for row_index, row in enumerate(table.rows):
                        for cell_index, cell in enumerate(row.cells):
                            
                            if "temas" in cell.text.lower():
                                cuadros_encontrados += 1
                                
                                texto_cuadros_debajo = []
                                for j in range(1, 4):  
                                    if row_index + j < len(table.rows):
                                        siguiente_fila = table.rows[row_index + j]
                                        texto_cuadro_debajo = siguiente_fila.cells[cell_index].text
                                        texto_cuadros_debajo.append(texto_cuadro_debajo)
                                resultados.append((i, texto_cuadros_debajo))

                

                return cuadros_encontrados, resultados

            cuadros_encontrados, resultados = buscar_temas_en_tabla(doc_path)

            contador = 2
            
            resultado_formateado = []
            for inicio, temas in resultados:
                for tema in temas:
                    resultado_formateado.append((contador, tema))
                    contador += 1


            def obtener_datos_tercera_columna_despues_de_subtema(doc_path):
                datos_tercera_columna = []

                palabra_clave_subtema = "subtema"

                ultima_tabla = None
                for table in doc.tables:
                    ultima_tabla = table

                if ultima_tabla:
                    posicion_subtema = None
                    for fila_index, fila in enumerate(ultima_tabla.rows):
                        for celda_index, celda in enumerate(fila.cells):
                            if palabra_clave_subtema.lower() in celda.text.lower():
                                posicion_subtema = (fila_index, celda_index)

                    if posicion_subtema:
                        fila_subtema, _ = posicion_subtema  
                        numero_filas = len(ultima_tabla.rows)

                        for fila in range(fila_subtema + 1, numero_filas):
                            if fila < numero_filas:  
                                celda = ultima_tabla.rows[fila].cells[2]
                                datos_celda = celda.text.strip()
                                datos_tercera_columna.append((fila, datos_celda))

                return datos_tercera_columna, ultima_tabla



            
            datos_tercera_columna_despues_subtemas, ultima_tabla = obtener_datos_tercera_columna_despues_de_subtema(doc_path)

            datos_tercera_columna_despues_subtema = resultado_formateado
            
            for fila, dato in datos_tercera_columna_despues_subtema:
                nuevo_dato = dato 
                celda_editar = ultima_tabla.rows[fila].cells[2]
                celda_editar.text = nuevo_dato
                for paragraph in ultima_tabla.rows[fila].cells[2].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)

            def obtener_datos_cuarta_columna_despues_de_concepto(doc_path):
                datos_tercera_columna = []

                palabra_clave_concepto = "concepto"

                ultima_tabla = None
                for table in doc.tables:
                    ultima_tabla = table

                if ultima_tabla:
                    
                    posicion_concepto = None
                    for fila_index, fila in enumerate(ultima_tabla.rows):
                        celda = fila.cells[3]  
                        if palabra_clave_concepto.lower() in celda.text.lower():
                            posicion_concepto = (fila_index, 3)  
                            break  

                    
                    if posicion_concepto:
                        fila_concepto, columna_concepto = posicion_concepto
                        numero_filas = len(ultima_tabla.rows)

                        for fila in range(fila_concepto + 1, numero_filas):
                            celda = ultima_tabla.rows[fila].cells[2]  
                            datos_celda = celda.text.strip()
                            datos_tercera_columna.append((fila, datos_celda))  

                return datos_tercera_columna, ultima_tabla


            resulta, ultima_tabla = obtener_datos_cuarta_columna_despues_de_concepto(doc_path)

            

            datos_contados = len(resulta)

            
            datos_formateados = []
            indice_deseado = 2

            for indice, dato in enumerate(conceptos):
                datos_formateados.append((indice_deseado, dato))
                indice_deseado += 1

                
                if indice_deseado > datos_contados + 1:  
                    break


            for fila, dato in datos_formateados:
                nuevo_dato = dato 
                celda_editar = ultima_tabla.rows[fila].cells[3]  
                celda_editar.text = nuevo_dato
                for paragraph in ultima_tabla.rows[fila].cells[3].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)

            def obtener_celda_despues_de_conocimientos_generales(doc_path):
                palabra_clave_concepto = "conocimientos generales del profesor"
                celda_despues_conocimientos = None

                
                ultima_tabla = doc.tables[-1]  

                for fila_index, fila in enumerate(ultima_tabla.rows):
                    for celda_index, celda in enumerate(fila.cells):
                        if palabra_clave_concepto.lower() in celda.text.lower():
                            
                            if fila_index + 1 < len(ultima_tabla.rows):
                                celda_despues_conocimientos = ultima_tabla.rows[fila_index + 1].cells[0]  
                            break
                    if celda_despues_conocimientos is not None:
                        break

                return celda_despues_conocimientos, ultima_tabla

            
            celda_despues_conocimientos, ultima_tabla = obtener_celda_despues_de_conocimientos_generales(doc_path)

            
            if celda_despues_conocimientos is not None:
                nuevo_dato = conocimientos  
                celda_despues_conocimientos.text = nuevo_dato
                primer_run = celda_despues_conocimientos.paragraphs[0].runs[0]
                primer_run.font.name = 'Arial'

            
            
            def obtener_dato_columna(doc_path):

                datos_tercera_columna = []

                palabra_clave_subtema = "tema"

                ultima_tabla = None
                for table in doc.tables:
                    ultima_tabla = table

                if ultima_tabla:
                    posicion_subtema = None
                    for fila_index, fila in enumerate(ultima_tabla.rows):
                        for celda_index, celda in enumerate(fila.cells):
                            if palabra_clave_subtema.lower() in celda.text.lower():
                                posicion_subtema = (fila_index, celda_index)

                    if posicion_subtema:
                        fila_subtema, columna_subtema = posicion_subtema
                        numero_filas = len(ultima_tabla.rows)

                        for fila in range(fila_subtema + 1, numero_filas):
                            celda = ultima_tabla.rows[fila].cells[1]
                            datos_celda = celda.text.strip()
                            datos_tercera_columna.append((fila, datos_celda))

                return datos_tercera_columna, ultima_tabla

            datos_tercera_columna_despues_subtemas, ultima_tabla = obtener_dato_columna(doc_path)

            datos_dict = {}

            
            for tupla in datos_tercera_columna_despues_subtemas:
                datos_dict[tupla[1]] = tupla

            
            datos_tercera_columna_despues_subtemas_sin_repetidos = list(datos_dict.values())

            
            temas = temas_a_usar
            datos_contados = len(datos_tercera_columna_despues_subtemas_sin_repetidos)

            datos_formateados = []

            
            for fila, (indice, dato) in enumerate(datos_tercera_columna_despues_subtemas_sin_repetidos, start=2):
                
                if fila - 2 < len(temas):
                    nuevo_dato = temas[fila - 2]
                else:
                    nuevo_dato = ""
                
                
                datos_formateados.append((indice, nuevo_dato))

            print(datos_formateados)
            
            for fila, dato in datos_formateados:
                nuevo_dato = dato
                celda_editar = ultima_tabla.rows[fila].cells[1]
                celda_editar.text = nuevo_dato
                for paragraph in ultima_tabla.rows[fila].cells[1].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
                        run.bold = True


                def guardar_imagen_desde_base64(base64_string, ruta_completa_archivo):
                    
                    _, base64_data = base64_string.split(",", 1)

                    try:
                        
                        image_data = base64.b64decode(base64_data)

                        
                        imagen = Image.open(BytesIO(image_data))

                        
                        imagen.save(ruta_completa_archivo, 'PNG')

                        
                        print("Imagen eliminada del disco.")
                    except Exception as e:
                        print("Error al guardar la imagen:", e)

                nombre_archivo = "imagen_guardada.png"
                ruta_carpeta = "static/images/"  

                
                ruta_completa_archivo = os.path.join(ruta_carpeta, nombre_archivo)

                guardar_imagen_desde_base64(firma_base64, ruta_completa_archivo)

                
                tabla = None
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if current_user.fullname.lower() in cell.text.lower():
                                tabla = table
                                break
                        if tabla:
                            break
                    if tabla:
                        break
            imagen_agregada = False
            
            if tabla:
                
                for i, row in enumerate(tabla.rows):
                    for j, cell in enumerate(row.cells):
                        if current_user.fullname.lower() in cell.text.lower():
                            firma_cell = tabla.cell(i - 2, j)
                            firma_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  
                            firma_cell.paragraphs[0].alignment_vertical = WD_ALIGN_VERTICAL.CENTER  

                            
                            if not imagen_agregada:
                                firma_cell.paragraphs[0].add_run().add_picture(ruta_completa_archivo, width=Inches(1.0), height=Inches(0.5))
                                imagen_agregada = True  

                            break
                    else:
                        continue
                    break
                
                    print("Documento Word con firma insertada guardado:")
                else:
                    print("No se encontró la tabla o el nombre 'jose uriel saenz cuellar' en el documento.")
                os.remove(ruta_completa_archivo)

            ruta_archivo_modificado = os.path.join(carpeta_modificados)

            with open(ruta_archivo_modificado, "wb") as f:
                doc.save(f)

            current_datetime = datetime.today().strftime('%Y-%m-%d')
            print(name)
            cursor = db.connection.cursor()
            cursor.execute("SELECT idPlantilla FROM `plantillas` WHERE namePlantilla = %s;", (name,))
            result = cursor.fetchone()

            if result:
                id = int(result[0])
                print(id)
            else:
                print("No data found.")


            sql = "INSERT INTO planeaciones (name, fecha, img, idMaestro, idPlantilla) VALUES (%s, %s, %s, %s, %s);"
            datos = (filename_with_timestamp, current_datetime, name, current_user.id, id,)
            cursor = db.connection.cursor()
            cursor.execute(sql, datos)
            db.connection.commit()
            cursor.close()

        doc_path = os.path.join('static/plantillas/', name)

        buscar_porcentajes_y_reemplazar(doc_path)

        return redirect(url_for('planeacionesU'))  
    
    else:
        return protected()

@app.route('/planeacionesU')
@login_required
def planeacionesU():
    if current_user.user_type == "Maestro":
        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM `materias` ORDER BY idMateria DESC;")
        materia = cursor.fetchall()
        cursor = db.connection.cursor()
        cursor.execute("""
            SELECT plantillas.*, materias.nombre, users.Nombres, users.Apellidos
            FROM plantillas
            INNER JOIN materias ON plantillas.idMateria = materias.idMateria
            INNER JOIN users ON plantillas.idUser = users.idUser
            WHERE (idMaestro1 = %s OR idMaestro2 = %s OR idMaestro3 = %s)
            AND NOT EXISTS (
                SELECT 1
                FROM planeaciones
                WHERE planeaciones.idPlantilla = plantillas.idPlantilla
            )
            ORDER BY plantillas.idPlantilla DESC;
        """, (current_user.id, current_user.id, current_user.id))
        plantilla = cursor.fetchall()
        cursor = db.connection.cursor()
        cursor.execute("""
            SELECT planeaciones.*, users.Nombres, users.Apellidos, materias.nombre
            FROM planeaciones
            INNER JOIN plantillas ON planeaciones.idPlantilla = plantillas.idPlantilla
            INNER JOIN users ON plantillas.idUser = users.idUser
            INNER JOIN materias ON plantillas.idMateria = materias.idMateria
            WHERE (planeaciones.idMaestro = %s)
            ORDER BY idPlan DESC;
        """, (current_user.id,))
        plan = cursor.fetchall()
        cursor.close()
        print(plantilla)
               
        return render_template('user/Rplaneacion.html', materia=materia, plantilla=plantilla, plan=plan)
    
    else:
        return protected()

@app.route('/form', methods=['POST'])
@login_required
def form():
    if current_user.user_type == "Maestro":
        idPlantilla = request.form['txtid']
        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM `plantillas` WHERE idPlantilla = %s;", (int(idPlantilla),))
        plantilla = cursor.fetchall()
        cursor.close()
        plantillas = plantilla

        def buscar_temas_en_tabla(doc_path):
            cuadros_encontrados = 0
            resultados = []

            doc = docx.Document(doc_path)

            for i, table in enumerate(doc.tables, 1):
                for row_index, row in enumerate(table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        
                        if "temas" in cell.text.lower():
                            cuadros_encontrados += 1
                            
                            texto_cuadros_debajo = []
                            for j in range(1, 4):  
                                if row_index + j < len(table.rows):
                                    siguiente_fila = table.rows[row_index + j]
                                    texto_cuadro_debajo = siguiente_fila.cells[cell_index].text
                                    texto_cuadros_debajo.append(texto_cuadro_debajo)
                            resultados.append((i, texto_cuadros_debajo))

            
            doc.save(doc_path)

            return cuadros_encontrados, resultados

        doc_path = 'static/plantillas/' + plantillas[0][1]

        cuadros_encontrados, resultados = buscar_temas_en_tabla(doc_path)
        resultados = resultados
        result = resultados
        
        nombres_temas = [tema for _, temas in resultados for tema in temas]
        


        def obtener_datos_columna_4_desde_fila_3(doc_path):
            
            doc = docx.Document(doc_path)

            
            ultima_tabla = None
            for table in doc.tables:
                ultima_tabla = table

            if ultima_tabla:
                
                datos_columna_4 = []

                for fila in ultima_tabla.rows[2:]:  
                    if 3 < len(fila.cells):  
                        celda_columna_4 = fila.cells[3]
                        texto_columna_4 = celda_columna_4.text.strip()
                        datos_columna_4.append(texto_columna_4)

                return datos_columna_4
            else:
                return None

        
        datos_columna_4 = obtener_datos_columna_4_desde_fila_3(doc_path)
        
        
        def buscar_temas_en_tabla(doc_path):
            cuadros_encontrados = 0
            resultados = []

            doc = docx.Document(doc_path)

            unidad_tematica = None  

            for i, table in enumerate(doc.tables, 1):
                for row_index, row in enumerate(table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        if cell_index + 1 < len(row.cells):
                            celda_derecha = row.cells[cell_index + 1]
                            if "unidad temática" in cell.text.lower():
                                unidad_tematica = celda_derecha
                                break
                    
                for row_index, row in enumerate(table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        if "temas" in cell.text.lower():
                            cuadros_encontrados += 1

                            if unidad_tematica:
                                texto_cuadros_debajo = []
                                for j in range(1, 4):
                                    if row_index + j < len(table.rows):
                                        siguiente_fila = table.rows[row_index + j]
                                        texto_cuadro_debajo = siguiente_fila.cells[cell_index].text
                                        texto_cuadros_debajo.append(texto_cuadro_debajo)
                                resultados.append((unidad_tematica.text, i, texto_cuadros_debajo))

            return cuadros_encontrados, resultados

        cuadros_encontrados, resulta = buscar_temas_en_tabla(doc_path)
        print(resulta)
        

        return render_template('user/form.html', temas=nombres_temas, plantilla=plantilla, datos=datos_columna_4, result=result, resulta=resulta)

    else:
        return protected()
    
@app.route('/download_file', methods=['POST'])
@login_required
def download_file():
    if current_user.user_type == "Administrador":
        archivo = request.form['txtarchivo']
        def word_to_pdf(input_file, output_file):
            
            if not input_file.lower().endswith('.docx'):
                raise ValueError("El archivo de entrada debe ser .docx")

            
            if not output_file.lower().endswith('.pdf'):
                raise ValueError("El archivo de salida debe ser .pdf")

            try:
                
                convert(input_file, output_file)

                print("Conversión completada: {} -> {}".format(input_file, output_file))

                
                temp_dir = os.path.join(os.path.expanduser("~"), "Downloads")
                temp_file = os.path.join(temp_dir, os.path.basename(output_file))
                shutil.move(output_file, temp_file)

                
                webbrowser.open(temp_file)


                time.sleep(10)
                os.remove(temp_file)
            except Exception as e:
                print("Error durante la conversión:", e)
            
        archivo_word = "static/planeacion/" + archivo

        
        archivo_pdf = "static/planeacion/archivo.pdf"

        word_to_pdf(archivo_word, archivo_pdf)
        
        return redirect(url_for('plan'))
    
    elif current_user.user_type == "Maestro":
        archivo = request.form['txtarchivo']
        def word_to_pdf(input_file, output_file):
            
            if not input_file.lower().endswith('.docx'):
                raise ValueError("El archivo de entrada debe ser .docx")

            
            if not output_file.lower().endswith('.pdf'):
                raise ValueError("El archivo de salida debe ser .pdf")

            try:
                
                convert(input_file, output_file)

                print("Conversión completada: {} -> {}".format(input_file, output_file))

                
                temp_dir = os.path.join(os.path.expanduser("~"), "Downloads")
                temp_file = os.path.join(temp_dir, os.path.basename(output_file))
                shutil.move(output_file, temp_file)

                
                webbrowser.open(temp_file)


                time.sleep(10)
                os.remove(temp_file)
            except Exception as e:
                print("Error durante la conversión:", e)
            
        archivo_word = "static/planeacion/" + archivo

        
        archivo_pdf = "static/planeacion/archivo.pdf"

        word_to_pdf(archivo_word, archivo_pdf)
        
        return redirect(url_for('planeacionesU'))
    
    else:
        protected()

@app.route('/plantillas')
@login_required
def plantillas():
    if current_user.user_type == "Administrador" and session['ass'] == 'plan':
        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM `materias` ORDER BY idMateria DESC;")
        materia = cursor.fetchall()
        cursor = db.connection.cursor()
        cursor.execute("""SELECT ma.* 
                       FROM maestros AS ma
                        JOIN grupos g ON ma.idGrupo = g.idGrupo
                        JOIN grado gr ON g.idGrado = gr.idGrado
                        JOIN areas a ON g.idArea = a.idArea
                        JOIN carrera c ON a.idCarrera = c.idCarrera
                        JOIN users u ON c.idCarrera = u.idCarrera
                       WHERE u.idUser = %s AND gr.nameGrado = %s AND a.idArea = %s 
                       ORDER BY idMaestro DESC; """, (current_user.id, session.get('grado'), session.get('area'),))
        maestro = cursor.fetchall()
        cursor = db.connection.cursor()
        cursor.execute("""SELECT plantillas.*, materias.nombre
        FROM plantillas
        INNER JOIN materias ON plantillas.idMateria = materias.idMateria
        JOIN maestros ms ON materias.idMaestro = ms.idMaestro
        JOIN grupos g ON ms.idGrupo = g.idGrupo
        JOIN grado gr ON g.idGrado = gr.idGrado
        JOIN areas a ON g.idArea = a.idArea
        JOIN carrera c ON a.idCarrera = c.idCarrera
        JOIN users u ON c.idCarrera = u.idCarrera
        WHERE plantillas.idUser = %s AND gr.nameGrado = %s AND a.idArea = %s
        ORDER BY plantillas.idPlantilla DESC;
        """,  (current_user.id, session.get('grado'), session.get('area'),))
        plantilla = cursor.fetchall()
        cursor.close()
        return render_template('admin/plantillas.html', materia=materia, maestro=maestro, plantilla=plantilla)
    
    else:
        return protected()

@app.route('/materia')
@login_required
def materia():
    if current_user.user_type == "Administrador" and session['ass'] == 'plan':
        cursor = db.connection.cursor()
        cursor.execute("""
            SELECT a.idArea, a.aName
            FROM users AS u
            INNER JOIN carrera AS c ON u.idCarrera = c.idCarrera
            INNER JOIN areas AS a ON c.idCarrera = a.idCarrera
            INNER JOIN grado AS g ON c.idGrado = g.idGrado
            WHERE u.idUser = %s AND g.nameGrado = %s
        """, (int(current_user.id), session.get('grado'),))
        special = cursor.fetchall()

        cursor = db.connection.cursor()
        cursor.execute("""SELECT g.*, a.aName, gr.nameGrado AS nombreGrado
        FROM grupos AS g
        INNER JOIN areas AS a ON g.idArea = a.idArea
        INNER JOIN carrera AS c ON a.idCarrera = c.idCarrera
        INNER JOIN users AS u ON c.idCarrera = u.idCarrera
        INNER JOIN grado AS gr ON g.idGrado = gr.idGrado
        WHERE u.idUser = %s AND gr.nameGrado = %s AND a.idArea = %s
        ORDER BY g.idGrupo DESC;
        """, (current_user.id, session.get('grado'), session.get('area'),))
        grupo = cursor.fetchall()
        
        cursor = db.connection.cursor()
        cursor.execute("""SELECT ma.* 
                       FROM maestros AS ma
                        JOIN grupos g ON ma.idGrupo = g.idGrupo
                        JOIN grado gr ON g.idGrado = gr.idGrado
                        JOIN areas a ON g.idArea = a.idArea
                        JOIN carrera c ON a.idCarrera = c.idCarrera
                        JOIN users u ON c.idCarrera = u.idCarrera
                       WHERE u.idUser = %s AND gr.nameGrado = %s AND a.idArea = %s 
                       ORDER BY idMaestro DESC; """, (current_user.id, session.get('grado'), session.get('area'),))
        maestros = cursor.fetchall()

        cursor = db.connection.cursor()
        cursor.execute("""SELECT m.*, a.aName AS nombre_area
            FROM materias m
            JOIN maestros ms ON m.idMaestro = ms.idMaestro 
            JOIN grupos g ON ms.idGrupo = g.idGrupo
            JOIN grado gr ON g.idGrado = gr.idGrado
            JOIN areas a ON g.idArea = a.idArea
            JOIN carrera c ON a.idCarrera = c.idCarrera
            JOIN users u ON c.idCarrera = u.idCarrera
            WHERE u.idUser = %s AND gr.nameGrado = %s AND a.idArea = %s
            ORDER BY m.idMateria DESC;
            """, (current_user.id, session.get('grado'), session.get('area'),))
        materia = cursor.fetchall()
        cursor.close()

        
        return render_template('admin/materias.html', area=special, materia=materia, maestros=maestros, grupo=grupo)
    
    else:
        return protected()
    
@app.route('/protected')
@login_required
def protected():
    return render_template('admin/protected.html')

@app.route('/protected_dica')
@login_required
def protected_dica():
    return render_template('admin/protected_dica.html')

@app.route('/plantilla_user')
@login_required
def plantilla_user():
    if current_user.user_type == "Maestro":
        return render_template('user/Rplaneacion.html')
    else:
        return protected()
    
@app.route('/plan')
@login_required
def plan():
    if current_user.user_type == "Administrador" and session['ass'] == 'plan':
        cursor = db.connection.cursor()
        cursor.execute("""
            SELECT pl.*, ma.nombre , m.Nombres, m.ApellidosP
            FROM planeaciones AS pl
            JOIN plantillas AS p ON p.idPlantilla = pl.idPlantilla
            JOIN materias AS ma ON p.idMateria = ma.idMateria
            JOIN maestros AS m ON pl.idMaestro = m.idMaestro
            JOIN grupos g ON m.idGrupo = g.idGrupo
            JOIN grado gr ON g.idGrado = gr.idGrado
            JOIN areas a ON g.idArea = a.idArea
            JOIN carrera c ON a.idCarrera = c.idCarrera
            JOIN users u ON c.idCarrera = u.idCarrera
            WHERE p.idUser = %s AND gr.nameGrado = %s AND a.idArea = %s
            ORDER BY pl.idPlan DESC;
        """,  (current_user.id, session.get('grado'), session.get('area'),))
        plan = cursor.fetchall()
        print(plan)

        return render_template('admin/planeacion.html', plan=plan)
    else:
        return protected()

@app.route('/inicio')
@login_required
def inicio():
    if current_user.user_type == 'Administrador' or current_user.user_type == 'Maestro' :
        return render_template('index.html')
    else:
        return protected()

    
@app.route('/inicio_dica')
@login_required
def inicio_dica():
    if current_user.user_type == 'Administrador':
        cursor = db.connection.cursor()
        cursor.execute("SELECT areas.aName, carrera.carreraName, areas.idArea FROM users INNER JOIN areas ON users.idCarrera = areas.idCarrera INNER JOIN carrera ON users.idCarrera = carrera.idCarrera WHERE users.idUser = %s ORDER BY users.idUser DESC;", (int(current_user.id),))
        area = cursor.fetchall()

        cursor.execute("SELECT DISTINCT grado.nameGrado FROM users INNER JOIN areas ON users.idCarrera = areas.idCarrera INNER JOIN carrera ON users.idCarrera = carrera.idCarrera INNER JOIN grado ON carrera.idGrado = grado.idGrado WHERE users.idUser = %s ORDER BY users.idUser DESC;", (int(current_user.id),))
        grado = cursor.fetchall()

        cursor.execute("SELECT carrera.carreraName FROM users INNER JOIN carrera ON users.idCarrera = carrera.idCarrera WHERE users.idUser = %s ORDER BY users.idUser DESC;", (int(current_user.id),))
        carrera = cursor.fetchall()

        carreras = carrera[0]
        grados = grado[0]

        return render_template('dica/index.html', area=area, grados=grados, carreras=carreras)  
     
    elif current_user.user_type == 'Maestro':
        cursor = db.connection.cursor()
        cursor.execute("""SELECT maestros.*, areas.aName, carrera.carreraName, areas.idArea
                       FROM maestros
                       INNER JOIN users ON users.idUser = maestros.idUser 
                       INNER JOIN areas ON users.idCarrera = areas.idCarrera 
                       INNER JOIN carrera ON users.idCarrera = carrera.idCarrera 
                       WHERE maestros.idMaestro = %s ORDER BY maestros.idMaestro DESC;""", (int(current_user.id),))
        area = cursor.fetchall()
        
        return render_template('dica/index.html', area=area)
    
    else:
        return protected()
    
@app.route('/inicio_plan')
@login_required
def inicio_plan():
    if current_user.user_type == 'Administrador':
        cursor = db.connection.cursor()
        cursor.execute("SELECT areas.aName, carrera.carreraName, areas.idArea FROM users INNER JOIN areas ON users.idCarrera = areas.idCarrera INNER JOIN carrera ON users.idCarrera = carrera.idCarrera WHERE users.idUser = %s ORDER BY users.idUser DESC;", (int(current_user.id),))
        area = cursor.fetchall()

        cursor.execute("SELECT DISTINCT grado.nameGrado FROM users INNER JOIN areas ON users.idCarrera = areas.idCarrera INNER JOIN carrera ON users.idCarrera = carrera.idCarrera INNER JOIN grado ON carrera.idGrado = grado.idGrado WHERE users.idUser = %s ORDER BY users.idUser DESC;", (int(current_user.id),))
        grado = cursor.fetchall()

        cursor.execute("SELECT carrera.carreraName FROM users INNER JOIN carrera ON users.idCarrera = carrera.idCarrera WHERE users.idUser = %s ORDER BY users.idUser DESC;", (int(current_user.id),))
        carrera = cursor.fetchall()

        carreras = carrera[0]
        grados = grado[0]

        return render_template('home.html', area=area, grados=grados, carreras=carreras)  
     
    elif current_user.user_type == 'Maestro':
        cursor = db.connection.cursor()
        cursor.execute("""SELECT materias.nombre, materias.idMateria
                       FROM materias 
                       INNER JOIN maestros ON materias.idMaestro = maestros.idMaestro
                       INNER JOIN grupos ON maestros.idGrupo = grupos.idGrupo
                       INNER JOIN grado ON grupos.idGrado = grado.idGrado
                       WHERE grado.nameGrado = 'Ingenieria' AND materias.idMaestro = %s;""", (current_user.id,))
        materia = cursor.fetchall()
        cursor.execute("""SELECT materias.nombre, materias.idMateria
                       FROM materias 
                       INNER JOIN maestros ON materias.idMaestro = maestros.idMaestro
                       INNER JOIN grupos ON maestros.idGrupo = grupos.idGrupo
                       INNER JOIN grado ON grupos.idGrado = grado.idGrado
                       WHERE grado.nameGrado = 'TSU';""")
        assig = cursor.fetchall()
        return render_template('home.html', materia=materia, assig=assig)
    
    else:
        return protected()

@app.route('/admins')
@login_required
def admins():
    if current_user.user_type == "Superusuario":

        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM carrera")
        carrera = cursor.fetchall()

        cursor.execute("SELECT users.*, carrera.carreraName FROM users INNER JOIN carrera ON users.idCarrera = carrera.idCarrera ORDER BY users.idUser DESC;")
        admin = cursor.fetchall()

        cursor.close()
        return render_template('superAdmin/admins.html', admin=admin, carrera=carrera)
    else:
        return protected()
    
@app.route('/activity')
@login_required
def activity():
    if current_user.user_type == "Maestro":
        resultr = []
        conte = []
        cursor = db.connection.cursor()
        cursor.execute("""SELECT students.* , g.cuatrimestre, g.nombre, gr.nameGrado, g.idGrupo
                       FROM students 
                       INNER JOIN grupos g ON students.idGrupo = g.idGrupo
                       INNER JOIN grado gr ON g.idGrado = gr.idGrado
                       WHERE students.idGrupo = %s ORDER BY idStudent DESC;""", (session['idMateria'],))
        student = cursor.fetchall()

        cursor = db.connection.cursor()
        cursor.execute("""SELECT m.*, p.name
            FROM materias m
            LEFT JOIN plantillas pl ON m.idMateria = pl.idMateria
            LEFT JOIN planeaciones p ON pl.idPlantilla = p.idPlantilla
            JOIN maestros ms ON m.idMaestro = ms.idMaestro 
            JOIN grupos g ON ms.idGrupo = g.idGrupo
            JOIN grado gr ON g.idGrado = gr.idGrado
            JOIN areas a ON g.idArea = a.idArea
            JOIN carrera c ON a.idCarrera = c.idCarrera
            JOIN users u ON c.idCarrera = u.idCarrera
            WHERE m.idMaestro = %s AND gr.nameGrado = %s AND a.idArea = %s AND pl.idMateria = %s
            ORDER BY m.idMateria DESC;
            """, (current_user.id, session.get('grado'), session.get('area'), session['id']))
        plan = cursor.fetchall()
        if session['unidad'] != 'promedio':
            cursor = db.connection.cursor()
            cursor.execute("SELECT * FROM dica WHERE unidad = %s AND idMateria = %s;", (session['unidad'], session['id']))
            dica = cursor.fetchall()

        if plan:
            name = plan[0][4]
            name = 'static/planeacion/' + name
        else:
            name = None
        cursor.close()
        grouped_values = {}
        
        
        
        if name != None:
            
            def buscar_temas_en_tabla(doc_path):
                cuadros_encontrados = 0
                resultados = []

                doc = docx.Document(doc_path)

                unidad_tematica = None  

                for i, table in enumerate(doc.tables, 1):
                    for row_index, row in enumerate(table.rows):
                        for cell_index, cell in enumerate(row.cells):
                            if cell_index + 1 < len(row.cells):
                                celda_derecha = row.cells[cell_index + 1]
                                if "unidad temática" in cell.text.lower():
                                    unidad_tematica = celda_derecha
                                    break
                        
                    for row_index, row in enumerate(table.rows):
                        for cell_index, cell in enumerate(row.cells):
                            if "temas" in cell.text.lower():
                                cuadros_encontrados += 1

                                if unidad_tematica:
                                    texto_cuadros_debajo = []
                                    for j in range(1, 4):
                                        if row_index + j < len(table.rows):
                                            siguiente_fila = table.rows[row_index + j]
                                            texto_cuadro_debajo = siguiente_fila.cells[cell_index].text
                                            texto_cuadros_debajo.append(texto_cuadro_debajo)
                                    resultados.append((unidad_tematica.text, i, texto_cuadros_debajo))

                return cuadros_encontrados, resultados
            
            
            def obtener_datos_tabla(doc_path):
                doc = docx.Document(doc_path)
                datos = []

                for table in doc.tables:
                    descripcion_columna = None
                    porcentaje_columna = None

                    for row_index, row in enumerate(table.rows):
                        for cell_index, cell in enumerate(row.cells):
                            texto = cell.text.lower()
                            
                            if "descripción de la actividad para alcanzar el resultado de aprendizaje de la unidad" in texto:
                                descripcion_columna = cell_index
                            elif "% de la evaluación de la unidad" in texto:
                                porcentaje_columna = cell_index

                        if descripcion_columna is not None and porcentaje_columna is not None:
                            break

                    if descripcion_columna is not None and porcentaje_columna is not None:
                        for row in table.rows[row_index + 1:]:  
                            descripcion = row.cells[descripcion_columna].text
                            porcentaje = row.cells[porcentaje_columna].text
                            if descripcion.strip() and porcentaje.strip():  
                                datos.append((descripcion, porcentaje))

                        break  

                return datos
            
            doc_path = name
            datos_tabla = obtener_datos_tabla(doc_path)
            ra_list = []
            examen_list = []
            ser_list = []
            otros_list = []

            for item in datos_tabla:
                if 'RA' in item[0]:
                    ra_list.append((item[0],))
                elif 'Examen' in item[0]:
                    examen_list.append((item[0],))
                elif 'Ser' in item[0]:
                    ser_list.append((item[0],))
                else:
                    otros_list.append((item[0],))


            formatted_lists = [len(ra_list), len(examen_list), 4, len(otros_list)]
            
            cuadros_encontrados, resultr = buscar_temas_en_tabla(doc_path)
            conte = len(resultr)
            ultra_dica = []
            final_data = []
            final_final_final = []
            if session['unidad'] == 'promedio':
                for intera in range(1, conte + 1) :
                    cursor = db.connection.cursor()
                    cursor.execute("SELECT * FROM dica WHERE unidad =  %s AND idMateria =%s;", (int(intera),session['id'],))
                    dica = cursor.fetchall()
                    cursor.close()
                    print("MYSONG")
                    def process_data(data, formatted_lists, is_actividad):
                        grouped_values = {}

                        for tup in data:
                            _, value, key, *_ = tup
                            if ("Actividad" in tup and is_actividad) or ("RA" in tup and not is_actividad):
                                if key in grouped_values:
                                    grouped_values[key].append((tup[-1], value))
                                else:
                                    grouped_values[key] = [(tup[-1], value)]

                        result_list = []

                        for key in grouped_values:
                            values_dict = {}
                            for last_number, value in grouped_values[key]:
                                values_dict[last_number] = value

                            values_to_add = [values_dict.get(i, 0) for i in range(1, formatted_lists[3 if is_actividad else 0] + 1)]
                            result_list.append([key] + values_to_add)

                        combined_list = []

                        for tup in student:
                            found = False
                            for sublista in result_list:
                                if tup[0] == sublista[0]:
                                    combined_tup = tup + (sublista,)
                                    combined_list.append(combined_tup)
                                    found = True
                                    break

                            if not found:
                                zeros_to_add = (0,) * (formatted_lists[3 if is_actividad else 0])
                                combined_list.append(tup + ((tup[0],) + zeros_to_add,))

                        for st in combined_list:
                            data_A = st[10]  
                            AB4 = len(data_A)  

                            if AB4 <= len(data_A) and AB4 > 1:
                                data_A = data_A[1:]  
                                result = sum(data_A) / (AB4 - 1)  
                            else:
                                result = data_A[0]

                            st_with_result = st + (result,)  
                            st_index = combined_list.index(st)  
                            combined_list[st_index] = st_with_result  

                        return combined_list

                    
                    student_with_actividad = process_data(dica, formatted_lists, is_actividad=True)
                    student_with_ra = process_data(dica, formatted_lists, is_actividad=False)
                    
                    def process_data_ser(data, formatted_lists, is_actividad):
                        grouped_values = {}

                        for tup in data:
                            _, value, key, *_ = tup
                            if ("Examen" in tup and is_actividad) or ("Ser" in tup and not is_actividad):
                                if key in grouped_values:
                                    grouped_values[key].append((tup[-1], value))
                                else:
                                    grouped_values[key] = [(tup[-1], value)]

                        result_list = []

                        for key in grouped_values:
                            values_dict = {}
                            for last_number, value in grouped_values[key]:
                                values_dict[last_number] = value

                            values_to_add = [values_dict.get(i, 0) for i in range(1, formatted_lists[1 if is_actividad else 2] + 1)]
                            result_list.append([key] + values_to_add)

                        combined_list = []

                        for tup in student:
                            found = False
                            for sublista in result_list:
                                if tup[0] == sublista[0]:
                                    combined_tup = tup + (sublista,)
                                    combined_list.append(combined_tup)
                                    found = True
                                    break

                            if not found:
                                zeros_to_add = (0,) * (formatted_lists[1 if is_actividad else 2])
                                combined_list.append(tup + ((tup[0],) + zeros_to_add,))

                        for st in combined_list:
                            data_A = st[10]  
                            AB4 = len(data_A)  

                            if AB4 <= len(data_A) and AB4 > 1:
                                data_A = data_A[1:]  
                                result = sum(data_A) / (AB4 - 1)  
                            else:
                                result = data_A[0]

                            st_with_result = st + (result,)  
                            st_index = combined_list.index(st)  
                            combined_list[st_index] = st_with_result  

                        return combined_list

                    
                    student_with_ex = process_data_ser(dica, formatted_lists, is_actividad=True)
                    student_with_ser = process_data_ser(dica, formatted_lists, is_actividad=False)
                    result_ubit = []
                    i = 0
                    for assig in student_with_actividad:
                        porcentaje_a = (assig[11] / 10) * 20
                        porcentaje_r = (student_with_ra[i][11] / 10) * 40
                        porcentaje_s = (student_with_ser[i][11] / 10) * 20
                        porcentaje_e = (student_with_ex[i][11] / 10) * 20
                        combined_tup = (*assig, ((porcentaje_e + porcentaje_s + porcentaje_a + porcentaje_r)/10))  
                        result_ubit.append(combined_tup)
                        i = i + 1

                    ultra_dica.append(result_ubit)
                
                data = ultra_dica
                unique_tuples = {}

                
                for sublist in data:
                    for tup in sublist:
                        id_key = tup[0]  
                        if id_key not in unique_tuples:
                            unique_tuples[id_key] = [tup]  
                        else:
                            unique_tuples[id_key].append(tup)  

                for id_key, tup_list in unique_tuples.items():
                    if len(tup_list) == 1:
                        final_data.extend(tup_list)  
                    else:
                        combined_tuple = tup_list[0][:-1] + ([tup[-1] for tup in tup_list],) + (sum(tup[-2] for tup in tup_list),) + (sum(tup[-1] for tup in tup_list),)
                        final_data.append(combined_tuple)  
                i = 0
                pr = 0
                dic=0
                final_final = []
                for assig in final_data:
                    pr = sum(assig[12])  # Suma de todas las notas
                    dic = len(assig[12])  # Cantidad de notas
                    promedio = pr / dic  # Calcular el promedio
                    porcentaje_a = promedio / 10 * 100  # Calcular el porcentaje
                    combined_tup = (*assig, round(promedio, 2))  # Crear una tupla combinada con promedio redondeado
                    final_final.append(combined_tup)
                final_final_final = final_final
                print(final_final)
            else:
                def process_data(data, formatted_lists, is_actividad):
                    grouped_values = {}

                    for tup in data:
                        _, value, key, *_ = tup
                        if ("Actividad" in tup and is_actividad) or ("RA" in tup and not is_actividad):
                            if key in grouped_values:
                                grouped_values[key].append((tup[-1], value))
                            else:
                                grouped_values[key] = [(tup[-1], value)]

                    result_list = []

                    for key in grouped_values:
                        values_dict = {}
                        for last_number, value in grouped_values[key]:
                            values_dict[last_number] = value

                        values_to_add = [values_dict.get(i, 0) for i in range(1, formatted_lists[3 if is_actividad else 0] + 1)]
                        result_list.append([key] + values_to_add)

                    combined_list = []

                    for tup in student:
                        found = False
                        for sublista in result_list:
                            if tup[0] == sublista[0]:
                                combined_tup = tup + (sublista,)
                                combined_list.append(combined_tup)
                                found = True
                                break

                        if not found:
                            zeros_to_add = (0,) * (formatted_lists[3 if is_actividad else 0])
                            combined_list.append(tup + ((tup[0],) + zeros_to_add,))

                    for st in combined_list:
                        data_A = st[10]  
                        AB4 = len(data_A)  

                        if AB4 <= len(data_A) and AB4 > 1:
                            data_A = data_A[1:]  
                            result = sum(data_A) / (AB4 - 1)  
                        else:
                            result = data_A[0]

                        st_with_result = st + (result,)  
                        st_index = combined_list.index(st)  
                        combined_list[st_index] = st_with_result  

                    return combined_list

                
                student_with_actividad = process_data(dica, formatted_lists, is_actividad=True)
                student_with_ra = process_data(dica, formatted_lists, is_actividad=False)
                print(student_with_actividad)
                print(student_with_ra)
                
                def process_data_ser(data, formatted_lists, is_actividad):
                    grouped_values = {}

                    for tup in data:
                        _, value, key, *_ = tup
                        if ("Examen" in tup and is_actividad) or ("Ser" in tup and not is_actividad):
                            if key in grouped_values:
                                grouped_values[key].append((tup[-1], value))
                            else:
                                grouped_values[key] = [(tup[-1], value)]

                    result_list = []

                    for key in grouped_values:
                        values_dict = {}
                        for last_number, value in grouped_values[key]:
                            values_dict[last_number] = value

                        values_to_add = [values_dict.get(i, 0) for i in range(1, formatted_lists[1 if is_actividad else 2] + 1)]
                        result_list.append([key] + values_to_add)

                    combined_list = []

                    for tup in student:
                        found = False
                        for sublista in result_list:
                            if tup[0] == sublista[0]:
                                combined_tup = tup + (sublista,)
                                combined_list.append(combined_tup)
                                found = True
                                break

                        if not found:
                            zeros_to_add = (0,) * (formatted_lists[1 if is_actividad else 2])
                            combined_list.append(tup + ((tup[0],) + zeros_to_add,))

                    for st in combined_list:
                        data_A = st[10]  
                        AB4 = len(data_A)  

                        if AB4 <= len(data_A) and AB4 > 1:
                            data_A = data_A[1:]  
                            result = sum(data_A) / (AB4 - 1)  
                        else:
                            result = data_A[0]

                        st_with_result = st + (result,)  
                        st_index = combined_list.index(st)  
                        combined_list[st_index] = st_with_result  

                    return combined_list

                
                student_with_ex = process_data_ser(dica, formatted_lists, is_actividad=True)
                student_with_ser = process_data_ser(dica, formatted_lists, is_actividad=False)
                print(student_with_ex)
                print(student_with_ser)
                result_ubit = []
                i = 0
                for assig in student_with_actividad:
                    porcentaje_a = (assig[11] / 10) * 20
                    porcentaje_r = (student_with_ra[i][11] / 10) * 40
                    porcentaje_s = (student_with_ser[i][11] / 10) * 20
                    porcentaje_e = (student_with_ex[i][11] / 10) * 20
                    combined_tup = (*assig, ((porcentaje_e + porcentaje_s + porcentaje_a + porcentaje_r)/10))  
                    result_ubit.append(combined_tup)
                    i = i + 1
                print(result_ubit)

            
            
            uni = session['unidad']
            return render_template('dica/activity.html', stu = student_with_ser, stude=student_with_ra ,stue=student_with_ex, student=student_with_actividad, dica=dica, conte=conte, uni=uni, formatted_lists=formatted_lists, result_ubit=result_ubit, final_data=final_final_final)
        else:
            return protected_dica()
    else:
        return protected()
    

@app.route('/config', methods=['POST'])
@login_required
def configuracion():
    if current_user.user_type == 'Maestro':
        idG = int(request.form['idG'])
        idM = session['id']
        idS = int(request.form['idS'])
        act = request.form['txtactividad']
        ra = request.form['txtra']
        reva = request.form['txtreva']
        unidad = request.form['txtunidad']
        integra = request.form['txtintegra']

        
        select_sql = "SELECT * FROM `dica` WHERE `idGrupo` = %s AND `idMaestro` = %s AND `idMateria` = %s;"
        select_data = (idG, idS, idM,)
        cursor = db.connection.cursor()
        cursor.execute(select_sql, select_data)
        existing_record = cursor.fetchone()

        if existing_record:
            
            update_sql = "UPDATE `dica` SET `actividades` = %s, `reva` = %s, `unidad` = %s, `integradora` = %s, `RA` = %s WHERE `idGrupo` = %s AND `idMaestro` = %s AND `idMateria` = %s;"
            update_data = (act, reva, unidad, integra, ra, idG, idS, idM, )
            cursor.execute(update_sql, update_data)
        else:
            
            insert_sql = "INSERT INTO `dica` (`actividades`, `reva`, `unidad`, `integradora`, `idGrupo`, `idMaestro`, `idMateria`, `RA`) VALUES (%s, %s, %s, %s, %s, %s, %s, %s);"
            insert_data = (act, reva, unidad, integra, idG, idS, idM, ra,)
            cursor.execute(insert_sql, insert_data)

        db.connection.commit()
        cursor.close()
        return redirect(url_for('activity'))

    else:
        return protected()
        
    
@app.route('/guardar', methods=['POST'])
def guardar_calificaciones():
    
    data_from_js = request.form
    tipo = data_from_js.get('tipo', '')
    student_data = {}
    print(tipo)
    for key, value in data_from_js.items():
        if key.startswith('txtcal_'):
            _, student_id, act = key.split('_')
            student_id = int(student_id)
            if student_id not in student_data:
                student_data[student_id] = []
            student_data[student_id].append((act, value))

    grouped_data = []

    for student_id, data in student_data.items():
        grouped_data.append((student_id, data))
    

    for grupo in grouped_data:
        student_id = grupo[0]
        for gru in grupo[1]:
            if gru[1] != '':
                select_sql = "SELECT * FROM `dica` WHERE `idStudent` = %s AND `token` = %s AND `tipo` = %s AND `unidad` = %s AND idMateria = %s;"
                select_data = (student_id, gru[0], tipo, session['unidad'], session['id'],)
                
                cursor = db.connection.cursor()
                cursor.execute(select_sql, select_data)
                existing_record = cursor.fetchone()
                print(existing_record)
                
                if existing_record:
                    
                    update_sql = "UPDATE `dica` SET `cali` = %s WHERE `idStudent` = %s AND `token` = %s AND `tipo` = %s AND `unidad` = %s AND idMateria = %s;"
                    update_data = (gru[1], student_id, gru[0], tipo, session['unidad'], session['id'],)
                    cursor.execute(update_sql, update_data)
                else:
                    
                    insert_sql = "INSERT INTO `dica` (`idStudent`, `cali`, `unidad`, `tipo`, `token`, `idMateria`) VALUES (%s, %s, %s, %s, %s, %s);"
                    insert_data = (student_id, gru[1], session['unidad'], tipo, gru[0], session['id'])
                    cursor.execute(insert_sql, insert_data)
                
                db.connection.commit()
                cursor.close()
    print(grouped_data)

    
    
    return redirect(url_for('activity'))  



@app.route('/promedio', methods=['POST'])
@login_required
def promedio():
    
    return redirect(url_for('activity'))


@app.route('/carreras')
@login_required
def carreras():
    if current_user.user_type == "Superusuario":
        cursor = db.connection.cursor()
        cursor.execute("SELECT carrera.*, grado.nameGrado, edificio.nameE FROM carrera INNER JOIN grado ON carrera.idGrado = grado.idGrado INNER JOIN edificio ON carrera.idEdificio = edificio.idEdificio ORDER BY idCarrera DESC;")
        career = cursor.fetchall()

        
        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM areas")
        scarrera = cursor.fetchall()

        cursor = db.connection.cursor()
        cursor.execute("SELECT idEdificio, nameE FROM edificio;")
        edificio = cursor.fetchall()

        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM `grado` ORDER BY idGrado DESC;")
        grado = cursor.fetchall()

        cursor.close()

        return render_template('superAdmin/careers.html', carrera=career, scarrera=scarrera, edificio=edificio, grado=grado)
    else:
        return protected()
    
@app.route('/especialidades')
@login_required
def especialidades():
    if current_user.user_type == "Superusuario":
        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM `carrera`;")
        career = cursor.fetchall()

        cursor = db.connection.cursor()
        cursor.execute("SELECT areas.*, carrera.carreraName FROM areas INNER JOIN carrera ON areas.idCarrera = carrera.idCarrera ORDER BY idArea DESC")
        special = cursor.fetchall()
        cursor.close()
        return render_template('superAdmin/specialties.html', career=career, special=special)
    else:
        return protected()
    
@app.route('/edificios')
@login_required
def edificios():
    if current_user.user_type == "Superusuario":
        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM `edificio` ORDER BY idEdificio DESC;")
        edificios = cursor.fetchall()
        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM `carrera`;")
        career = cursor.fetchall()
        cursor.close()
        return render_template('superAdmin/edificio.html', edificios=edificios,  carrera=career)
    else:
        return protected()
    
@app.route('/grado')
@login_required
def grado():
    if current_user.user_type == "Superusuario":
        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM `grado` ORDER BY idGrado DESC;")
        grado = cursor.fetchall()
        cursor.close()
        return render_template('superAdmin/grado.html', grado=grado)
    else:
        return protected()
    
@app.route('/gradoEdit', methods=['POST'])
@login_required
def gradoEdit():
    if current_user.user_type == "Superusuario":
        grado_id = request.form['id']
        nuevo_grado = request.form['newtxtnombre']
        image_url = get_image_url(nuevo_grado)
        cursor = db.connection.cursor()
        cursor.execute("UPDATE grado SET nameGrado = %s, img = %s WHERE idGrado = %s;", (nuevo_grado, image_url, grado_id))
        grado = cursor.fetchall()
        db.connection.commit()
        cursor.close()

        return redirect(url_for('grado'))
    else:
        return protected()

@app.route('/grupo')
@login_required
def grupo():
    if current_user.user_type == "Administrador" and session['ass'] == 'plan':
        cursor = db.connection.cursor()
        cursor.execute("""SELECT g.*, a.aName, gr.nameGrado AS nombreGrado
        FROM grupos AS g
        INNER JOIN areas AS a ON g.idArea = a.idArea
        INNER JOIN carrera AS c ON a.idCarrera = c.idCarrera
        INNER JOIN users AS u ON c.idCarrera = u.idCarrera
        INNER JOIN grado AS gr ON g.idGrado = gr.idGrado
        WHERE u.idUser = %s AND gr.nameGrado = %s AND a.idArea = %s
        ORDER BY g.idGrupo DESC;
        """, (current_user.id,  session.get('grado'),  session.get('area'),))

        grupo = cursor.fetchall()
        print(grupo)
        cursor = db.connection.cursor()
        cursor.execute("""
            SELECT a.idArea, a.aName
            FROM users AS u
            INNER JOIN carrera AS c ON u.idCarrera = c.idCarrera
            INNER JOIN areas AS a ON c.idCarrera = a.idCarrera
            INNER JOIN grado AS g ON g.idGrado = c.idGrado
            WHERE u.idUser = %s AND g.nameGrado = %s
        """, (current_user.id, session.get('grado'),))
        special = cursor.fetchall()

        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM grado WHERE nameGrado = %s", (session.get('grado'),))
        grados = cursor.fetchall()
        cursor.close()

        return render_template('admin/grupo.html', grupo=grupo, special=special, grados=grados)
    else:
        return protected()

@app.route('/alumnos')
@login_required
def alumnos():
    if current_user.user_type == "Administrador" and session['ass'] == 'dica':
        cursor = db.connection.cursor()
        cursor.execute("""SELECT g.*, a.aName, gr.nameGrado AS nombreGrado
        FROM grupos AS g
        INNER JOIN areas AS a ON g.idArea = a.idArea
        INNER JOIN carrera AS c ON a.idCarrera = c.idCarrera
        INNER JOIN users AS u ON c.idCarrera = u.idCarrera
        INNER JOIN grado AS gr ON g.idGrado = gr.idGrado
        WHERE u.idUser = %s AND gr.nameGrado = %s AND a.idArea = %s
        ORDER BY g.idGrupo DESC;
        """, (current_user.id,  session.get('grado'),  session.get('area'),))
        grupo = cursor.fetchall()
        cursor = db.connection.cursor()
        cursor.execute("""SELECT ma.* 
                       FROM students AS ma
                        JOIN grupos g ON ma.idGrupo = g.idGrupo
                        JOIN grado gr ON g.idGrado = gr.idGrado
                        JOIN areas a ON g.idArea = a.idArea
                        JOIN carrera c ON a.idCarrera = c.idCarrera
                        JOIN users u ON c.idCarrera = u.idCarrera
                       WHERE u.idUser = %s AND gr.nameGrado = %s AND a.idArea = %s 
                       ORDER BY idStudent DESC; """, (current_user.id, session.get('grado'), session.get('area'),))
        alumnos = cursor.fetchall()
        
        print(alumnos)
        cursor.close()

        return render_template('dica/admin/alumnos.html', grupo=grupo, alumnos=alumnos)
    else:
        return protected()

@app.route('/alumnosAdd', methods=['POST'])
@login_required
def alumnosAdd():
    if current_user.user_type == "Administrador":
       
        file = request.files['txtnombre']
        grupo = int(request.form['txtgrupo'])
        archivo_excel = file
        libro_trabajo = openpyxl.load_workbook(archivo_excel)
        hoja_activa = libro_trabajo.active
        datos = []

        for fila in hoja_activa.iter_rows(values_only=True):
            datos.append(fila)

        
        numero_tuplas = len(datos)

        
        print("Número de tuplas:", numero_tuplas)

        

        for tupla in datos:
            full_name = tupla[1]  
            matricula = tupla[0]  
            grupo = int(request.form['txtgrupo'])  

            partes_nombre = full_name.split()
            if len(partes_nombre) == 3:
                nombre_alumno = partes_nombre[0]
                apellido_paterno = partes_nombre[-2]
                apellido_materno = partes_nombre[-1]
            elif len(partes_nombre) == 4:
                nombre_alumno = " ".join(partes_nombre[0:2])
                apellido_paterno = partes_nombre[-2]
                apellido_materno = partes_nombre[-1]
            else:
                nombre_alumno = full_name
                apellido_paterno = ""
                apellido_materno = ""

            sql = "INSERT INTO `students` (`names`, `lastnameP`, `lastnameM`, `matricula`, `idGrupo`) VALUES (%s, %s, %s, %s, %s);"
            datos_insert = (nombre_alumno, apellido_paterno, apellido_materno, matricula, grupo)
            cursor = db.connection.cursor()
            cursor.execute(sql, datos_insert)
            db.connection.commit()
            cursor.close()
       
        return redirect(url_for('alumnos'))
    else:
        return protected()



def generar_contrasena_temporal():
    longitud = 10
    caracteres = string.ascii_letters + string.digits + string.punctuation
    contrasena = ''.join(random.choice(caracteres) for i in range(longitud))
    return contrasena

def enviar_correo(correo_destino, contrasena_temporal):
    
    smtp_server = "smtp.gmail.com"
    puerto_smtp = 587  
    usuario_smtp = "a70084681@gmail.com"
    contrasena_smtp = "akoualcntncsalhv"
    
    mensaje = MIMEMultipart()
    mensaje["From"] = usuario_smtp
    mensaje["To"] = correo_destino
    mensaje["Subject"] = "Restablecimiento de contraseña"

    contenido = f"Hola,\n\nHas solicitado restablecer tu contraseña. Porfavor introduce este token para cambiar tu contraseña: {contrasena_temporal}\n\nGracias."
    mensaje.attach(MIMEText(contenido, "plain"))

    try:
        server = smtplib.SMTP(smtp_server, puerto_smtp)
        server.starttls()
        server.login(usuario_smtp, contrasena_smtp)
        server.sendmail(usuario_smtp, correo_destino, mensaje.as_string())
        print("Correo electrónico enviado correctamente.")
        server.quit()
    except Exception as e:
        print("Error al enviar el correo electrónico:", str(e))

@app.route("/recuperar", methods=["GET", "POST"])
def recuperar():
    if request.method == "POST":
        correo = request.form["txtemail"]
        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM `users` WHERE correo=%s;", (correo,))
        usuario = cursor.fetchall()

        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM `maestros` WHERE correo=%s;", (correo,))
        maestros = cursor.fetchall()

        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM `superuser` WHERE correo=%s;", (correo,))
        super = cursor.fetchall()
        cursor.close()
        if usuario:
            contrasena_temporal = generar_contrasena_temporal()
            enviar_correo(correo, contrasena_temporal)  
            flash("Se ha enviado un Token temporal a su correo electrónico.")
            id = usuario[0][0]
            token = contrasena_temporal
            cursor = db.connection.cursor()
            cursor.execute("UPDATE users SET token=%s WHERE idUser = %s;", (token, id,))
            db.connection.commit()
            cursor.close()
            name="usuario"
            return render_template("auth/change.html", name=name)
        elif maestros:
            contrasena_temporal = generar_contrasena_temporal()
            enviar_correo(correo, contrasena_temporal)  
            flash("Se ha enviado un Token temporal a su correo electrónico.")
            id = usuario[0][0]
            token = contrasena_temporal
            cursor = db.connection.cursor()
            cursor.execute("UPDATE maestros SET token=%s WHERE idMaestro = %s;", (token, id,))
            db.connection.commit()
            cursor.close()
            name="maestros"
            return render_template("auth/change.html", name=name)
        elif super:
            contrasena_temporal = generar_contrasena_temporal()
            enviar_correo(correo, contrasena_temporal)  
            flash("Se ha enviado un Token temporal a su correo electrónico.")
            id = usuario[0][0]
            token = contrasena_temporal
            cursor = db.connection.cursor()
            cursor.execute("UPDATE superuser SET token=%s WHERE idsuser = %s;", (token, id,))
            db.connection.commit()
            cursor.close()
            name="super"
            return render_template("auth/change.html", name=name)
        else:
            flash("El correo electrónico no está registrado.", "error")
    return render_template("auth/recuperar.html")

@app.route("/change", methods=["POST"])
def change():
    if request.method == "POST":
        token = request.form["txttoken"]
        passw = request.form["txtpass"]
        name = request.form["txtname"]
        hashed_password = User.password(passw)
        tokens = ""
        error_message = "Error: el token no es válido."

        if name == "usuario":
            cursor = db.connection.cursor()
            cursor.execute("UPDATE users SET password=%s, token=%s WHERE token=%s;", (hashed_password, tokens, token,))
            db.connection.commit()
            cursor.close()
            if cursor.rowcount > 0:
                flash("Contraseña actualizada exitosamente.")
            else:
                flash(error_message)
                return render_template("auth/change.html")

        elif name == "maestros":
            cursor = db.connection.cursor()
            cursor.execute("UPDATE maestros SET password=%s, token=%s WHERE token=%s;", (hashed_password, tokens, token,))
            db.connection.commit()
            cursor.close()
            if cursor.rowcount > 0:
                flash("Contraseña actualizada exitosamente.")
            else:
                flash(error_message)
                return render_template("auth/change.html")

        elif name == "super":
            cursor = db.connection.cursor()
            cursor.execute("UPDATE superuser SET password=%s, token=%s WHERE token=%s;", (hashed_password, tokens, token,))
            db.connection.commit()
            cursor.close()
            if cursor.rowcount > 0:
                flash("Contraseña actualizada exitosamente.")
            else:
                flash(error_message)
                return render_template("auth/change.html")

    return render_template("auth/login.html")

@app.route('/edificio', methods=['POST'])
@login_required
def edificio():
    name = request.form['txtnombre']
    sql = "INSERT INTO `edificio` (`nameE`) VALUES (%s);"
    datos = (name,)
    cursor = db.connection.cursor()
    cursor.execute(sql, datos)
    db.connection.commit()
    cursor.close()
    
    return redirect(url_for('edificios'))

@app.route('/delete_edificio/<int:edificio_id>', methods=['GET'])
@login_required
def delete_edificio(edificio_id):

    cursor = db.connection.cursor()
    cursor.execute("DELETE FROM `edificio` WHERE `idEdificio` = %s;", (edificio_id,))
    db.connection.commit()
    cursor.close()

    return redirect(url_for('edificios'))

@app.route('/editEdificio/<int:idEdificio>', methods=['GET'])
@login_required
def editEdificio(idEdificio):
    cursor = db.connection.cursor()
    cursor.execute("SELECT * FROM `edificio` WHERE `idEdificio` = %s;", (idEdificio,))
    maestro_data = cursor.fetchone()
    cursor.close()
    if not maestro_data:
        return jsonify(), 404  
    column_names = [desc[0] for desc in cursor.description]
    maestro_dict = dict(zip(column_names, maestro_data))

    return jsonify(maestro_dict)

@app.route('/EdificioEdit', methods=['POST'])
@login_required
def EdificioEdit():
    if current_user.user_type == "Superusuario":
        edificio_id = request.form['id']
        nuevo_edificio = request.form['txtnombre']
        cursor = db.connection.cursor()
        cursor.execute("UPDATE edificio SET nameE = %s WHERE idEdificio = %s;", (nuevo_edificio, edificio_id))
        edificio = cursor.fetchall()
        db.connection.commit()
        cursor.close()

        return redirect(url_for('edificios'))
    else:
        return protected()

@app.route('/grados', methods=['POST'])
@login_required
def grados():
    name = request.form['txtnombre']
    image_url = get_image_url(name)
    sql = "INSERT INTO `grado` (`nameGrado`, `img`) VALUES (%s, %s);"
    datos = (name, image_url,)
    cursor = db.connection.cursor()
    cursor.execute(sql, datos)
    db.connection.commit()
    cursor.close()
    
    return redirect(url_for('grado'))

@app.route('/grados/<int:grados_id>', methods=['GET'])
@login_required
def delete_grados(grados_id):

    cursor = db.connection.cursor()
    cursor.execute("DELETE FROM `grado` WHERE `idGrado` = %s;", (grados_id,))
    db.connection.commit()
    cursor.close()

    return redirect(url_for('grado'))


@app.route('/admin', methods=['POST'])
@login_required
def admin():
    name = request.form['txtnombre']
    lastN = request.form['txtapellido']
    correo = request.form['txtcorreo']
    username = request.form['txtuser']
    number = request.form['txtnumero']
    passw = request.form['txtpasword']
    carrera = int(request.form['txtcarrera'])
    hashed_password = User.password(passw)

    sql = "INSERT INTO `users` (`Nombres`, `Apellidos`, `username`, `numero`, `password`, `idCarrera`, `correo`) VALUES (%s,%s,%s, %s,%s, %s,%s);"
    datos = (name, lastN, username, number, hashed_password, carrera, correo,)
    cursor = db.connection.cursor()
    cursor.execute(sql, datos)
    db.connection.commit()
    cursor.close()
    
    return redirect(url_for('admins'))

@app.route('/admin/<int:admin_id>', methods=['GET'])
@login_required
def delete_admin(admin_id):

    cursor = db.connection.cursor()
    cursor.execute("DELETE FROM `users` WHERE `idUser` = %s;", (admin_id,))
    db.connection.commit()
    cursor.close()

    return redirect(url_for('admins'))


@app.route('/editadmin/<int:iduser>', methods=['GET'])
@login_required
def editadmin(iduser):
    cursor = db.connection.cursor()
    cursor.execute("SELECT users.*, carrera.carreraName FROM users INNER JOIN carrera ON users.idCarrera = carrera.idCarrera WHERE idUser = %s;", (iduser,))
    grado_data = cursor.fetchone() 
    cursor.close()

    if not grado_data:
        return jsonify(), 404  
    column_names = [desc[0] for desc in cursor.description]
    grado_dict = dict(zip(column_names, grado_data))

    return jsonify(grado_dict)

@app.route('/adminedit', methods=['POST'])
@login_required
def adminedit():
    if current_user.user_type == "Superusuario":
        id = request.form['id']
        name = request.form['txtnombre']
        lastN = request.form['txtapellido']
        correo = request.form['txtcorreo']
        username = request.form['txtuser']
        number = request.form['txtnumero']
        carrera = int(request.form['txtcarrera'])
        if request.form['txtnewpasword'] == '':
            hashed_password = request.form['txtpasword']
        else:
            passw = request.form['txtnewpasword']
            hashed_password = User.password(passw)
        
        cursor = db.connection.cursor()
        cursor.execute("UPDATE users SET Nombres=%s, Apellidos=%s, username=%s, numero=%s, password=%s, idCarrera=%s, correo=%s WHERE idUser = %s;", (name, lastN, username, number, hashed_password, carrera, correo, id))
        carrera = cursor.fetchall()
        db.connection.commit()
        cursor.close()

        return redirect(url_for('admins'))
    else:
        return protected()



@app.route('/area', methods=['POST'])
@login_required
def area():
    name = request.form['txtnombre']
    carrer = int(request.form['txtcarrera'])
    image_url = get_image_url(name)

    sql = "INSERT INTO `areas` (`aName`, `idCarrera`, `urlA`) VALUES (%s, %s, %s);"
    datos = (name, carrer, image_url,)
    cursor = db.connection.cursor()
    cursor.execute(sql, datos)
    db.connection.commit()
    cursor.close()
    
    return redirect(url_for('especialidades'))

@app.route('/area/<int:area_id>', methods=['GET'])
@login_required
def delete_area(area_id):
    cursor = db.connection.cursor()
    cursor.execute("DELETE FROM `areas` WHERE `idArea` = %s;", (area_id,))
    db.connection.commit()
    cursor.close()


    return redirect(url_for('especialidades'))

@app.route('/editarea/<int:idarea>', methods=['GET'])
@login_required
def editarea(idarea):
    cursor = db.connection.cursor()
    cursor.execute("""SELECT areas.*, carrera.carreraName
    FROM areas
    INNER JOIN carrera ON areas.idCarrera = carrera.idCarrera
    WHERE areas.idArea = %s;""", (idarea,))
    grado_data = cursor.fetchone() 
    cursor.close()

    if not grado_data:
        return jsonify(), 404  
    column_names = [desc[0] for desc in cursor.description]
    grado_dict = dict(zip(column_names, grado_data))

    return jsonify(grado_dict)

@app.route('/areaEdit', methods=['POST'])
@login_required
def areaEdit():
    if current_user.user_type == "Superusuario":
        area_id = request.form['id']
        nuevo_area = request.form['txtnombre']
        nuevo_carrera = int(request.form['txtcarrera'])
        image_url = get_image_url(nuevo_area)
        cursor = db.connection.cursor()
        cursor.execute("UPDATE areas SET aName = %s, idCarrera = %s, urlA = %s WHERE idArea = %s;", (nuevo_area, nuevo_carrera, image_url, area_id))
        area = cursor.fetchall()
        db.connection.commit()
        cursor.close()

        return redirect(url_for('especialidades'))
    else:
        return protected()


@app.route('/carrera', methods=['POST'])
@login_required
def carrera():
    name = request.form['txtnombre']
    grado = int(request.form['txtGrado'])
    edificio = int(request.form['txtedificio'])
    image_url = get_image_url(name)

    sql = "INSERT INTO `carrera` (`carreraName`, `idedificio`, `img`, `idGrado`) VALUES (%s, %s, %s, %s);"
    datos = (name, edificio, image_url, grado,)
    cursor = db.connection.cursor()
    cursor.execute(sql, datos)
    db.connection.commit()
    cursor.close()
    
    return redirect(url_for('carreras'))

@app.route('/carrera/<int:idCarrera>', methods=['GET'])
@login_required
def delete_carrera(idCarrera):
    cursor = db.connection.cursor()
    cursor.execute("DELETE FROM `carrera` WHERE `idCarrera` = %s;", (idCarrera,))
    db.connection.commit()
    cursor.close()

    return redirect(url_for('carreras'))

@app.route('/editcarrera/<int:idcarrera>', methods=['GET'])
@login_required
def editcarrera(idcarrera):
    cursor = db.connection.cursor()
    cursor.execute("SELECT carrera.*, grado.nameGrado, edificio.nameE FROM carrera INNER JOIN grado ON carrera.idGrado = grado.idGrado INNER JOIN edificio ON carrera.idEdificio = edificio.idEdificio WHERE idCarrera = %s;", (idcarrera,))
    grado_data = cursor.fetchone() 
    cursor.close()

    if not grado_data:
        return jsonify(), 404  
    column_names = [desc[0] for desc in cursor.description]
    grado_dict = dict(zip(column_names, grado_data))

    return jsonify(grado_dict)

@app.route('/carreraedit', methods=['POST'])
@login_required
def carreraedit():
    if current_user.user_type == "Superusuario":
        carrera_id = request.form['id']
        name = request.form['txtnombre']
        grado = int(request.form['txtGrado'])
        edificio = int(request.form['txtedificio'])
        image_url = get_image_url(name)
        cursor = db.connection.cursor()
        cursor.execute("UPDATE carrera SET carreraName=%s, idGrado=%s, idedificio=%s, img=%s WHERE idCarrera = %s;", (name, grado, edificio, image_url, carrera_id))
        carrera = cursor.fetchall()
        db.connection.commit()
        cursor.close()

        return redirect(url_for('carreras'))
    else:
        return protected()


@app.route('/maestrosR', methods=['POST'])
@login_required
def maestrosR():
    name = request.form['txtnombre']
    lastNP = request.form['txtapellido']
    lastNM = request.form['txtapellidoM']
    correo = request.form['txtcorreo']
    user = request.form['txtusername']
    number = request.form['txtnumero']
    grupo = int(request.form['txtgrupo'])
    passw = request.form['txtpasword']
    hashed_password = User.password(passw)
    admin = int(current_user.id)
    
    sql = "INSERT INTO `maestros` (`Nombres`, `ApellidosP`, `ApellidoM`, `username`, `numero`, `password`, `correo`, `idUser`, `idGrupo`) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s);"
    datos = (name, lastNP, lastNM, user, number, hashed_password, correo, admin, grupo,)
    cursor = db.connection.cursor()
    cursor.execute(sql, datos)
    db.connection.commit()
    cursor.close()

    return redirect(url_for('register'))

@app.route('/maestroDelete/<int:maestro_id>', methods=['GET'])
@login_required
def delete_maestro(maestro_id):

    cursor = db.connection.cursor()
    cursor.execute("DELETE FROM `maestros` WHERE `idMaestro` = %s;", (maestro_id,))
    db.connection.commit()
    cursor.close()

    return redirect(url_for('register'))

@app.route('/maestroEdit', methods=['POST'])
@login_required
def maestrosEdit():
    name = request.form['txtnombre']
    lastNP = request.form['txtapellidoP']
    lastNM = request.form['txtapellidoM']
    correo = request.form['txtcorreo']
    user = request.form['txtusername']
    grupo = int(request.form['txtgrupo'])
    number = request.form['txtnumero']
    if request.form['txtpaswordnew'] == '':
        hashed_password = request.form['txtpasword']
    else:
        passw = request.form['txtpaswordnew']
        hashed_password = User.password(passw)
    id = request.form['id']

    sql = "UPDATE `maestros` SET `Nombres` = %s, `ApellidosP` = %s, `ApellidoM` = %s, `username` = %s, `numero` = %s, `password` = %s, `correo` = %s, `idGrupo` = %s WHERE `idMaestro` = %s ;"
    datos = (name, lastNP, lastNM, user, number, hashed_password, correo, grupo, id)
    cursor = db.connection.cursor()
    cursor.execute(sql, datos)
    db.connection.commit()
    cursor.close()

    return redirect(url_for('register'))

@app.route('/materiasAdd', methods=['POST'])
@login_required
def materiasAdd():
    name = request.form['txtnombre']
    maestro = int(request.form['txtmaestro'])
    image_url = get_image_url(name)
    sql = "INSERT INTO `materias` (`nombre`, `img`) VALUES (%s, %s);"
    datos = (name, image_url)
    cursor = db.connection.cursor()
    cursor.execute(sql, datos)
    db.connection.commit()
    cursor.close()
    
    return redirect(url_for('materia'))

@app.route('/delete_materia/<int:idmateria>', methods=['GET'])
@login_required
def delete_materia(idmateria):
    cursor = db.connection.cursor()
    cursor.execute("DELETE FROM `materias` WHERE `idMateria` = %s;", (idmateria,))
    db.connection.commit()
    cursor.close()

    return redirect(url_for('materia'))

@app.route('/editmateria/<int:idMateria>', methods=['GET'])
@login_required
def editmateria(idMateria):
    cursor = db.connection.cursor()
    cursor.execute("""SELECT materias.idMateria, materias.nombre AS nombre_materia, 
        materias.idMaestro, maestros.Nombres AS nombre_maestro, maestros.ApellidosP AS apellidos_maestro, maestros.ApellidoM AS apellidoM_maestro FROM materias
        INNER JOIN maestros ON materias.idMaestro = maestros.idMaestro
        WHERE materias.idMateria = %s;
    """, (idMateria,))
    materia_data = cursor.fetchone() 


    if not materia_data:
        return jsonify(), 404  
    column_names = [desc[0] for desc in cursor.description]
    materia_dict = dict(zip(column_names, materia_data))
    cursor.close()
    return jsonify(materia_dict)

@app.route('/materiasEdit', methods=['POST'])
@login_required
def materiasEdit():
    if current_user.user_type == "Administrador":
        id = request.form['id']
        name = request.form['txtnombre']
        maestro = int(request.form['txtmaestro'])
        image_url = get_image_url(name)
        
        cursor = db.connection.cursor()
        cursor.execute("UPDATE materias SET nombre=%s, img=%s, idMaestro=%s WHERE idMateria = %s;", (name, image_url, maestro, id,))
        materia = cursor.fetchall()
        db.connection.commit()
        cursor.close()

        return redirect(url_for('materia'))
    else:
        return protected()

@app.route('/alumnoEdit', methods=['POST'])
@login_required
def alumnoEdit():
    if current_user.user_type == "Administrador":
        idstudent = request.form['id']
        name = request.form['txtnombre']
        lastnameP = request.form['txtapellidoP']
        lastnameM = request.form['txtapellidoM']
        matricula = request.form['txtmatricula']
        grupo = int(request.form['txtgrupo'])
        cursor = db.connection.cursor()
        cursor.execute("UPDATE students SET names = %s, lastnameP = %s, lastnameM = %s, matricula = %s, idGrupo = %s WHERE idStudent = %s;", (name, lastnameP, lastnameM, matricula, grupo, idstudent))
        grado = cursor.fetchall()
        db.connection.commit()
        cursor.close()

        return redirect(url_for('alumnos'))
    else:
        return protected()

@app.route('/grupoA', methods=['POST'])
@login_required
def grupoA():
    name = request.form['txtnombre']
    idArea = int(request.form['txtArea'])
    cuatri = int(request.form['txtcuatri'])
    grado = int(request.form['txtGrado'])
    
    sql = "INSERT INTO `grupos` (`idGrupo`, `idArea`, `nombre`, `cuatrimestre`, `idGrado`) VALUES (NULL, %s, %s, %s, %s);"
    datos = (idArea, name, cuatri, grado,)
    cursor = db.connection.cursor()
    cursor.execute(sql, datos)
    db.connection.commit()
    cursor.close()

    return redirect(url_for('grupo'))

@app.route('/grupos/<int:idgrupos>', methods=['GET'])
@login_required
def delete_grupos(idgrupos):
    cursor = db.connection.cursor()
    cursor.execute("DELETE FROM `grupos` WHERE `idGrupo` = %s;", (idgrupos,))
    db.connection.commit()
    cursor.close()

    return redirect(url_for('grupo'))

@app.route('/editgrupo/<int:idGrupo>', methods=['GET'])
@login_required
def editgrupo(idGrupo):
    cursor = db.connection.cursor()
    cursor.execute("""SELECT g.*, a.aName AS nombreArea, gr.nameGrado AS nombreGrado
    FROM grupos AS g
    INNER JOIN areas AS a ON g.idArea = a.idArea
    INNER JOIN grado AS gr ON g.idGrado = gr.idGrado
    WHERE g.idGrupo = %s;
    """, (idGrupo,))
    grado_data = cursor.fetchone() 
    

    if not grado_data:
        return jsonify(), 404  
    column_names = [desc[0] for desc in cursor.description]
    grado_dict = dict(zip(column_names, grado_data))
    cursor.close()
    return jsonify(grado_dict)

@app.route('/grupoedit', methods=['POST'])
@login_required
def grupoedit():
    if current_user.user_type == "Administrador":
        id = request.form['id']
        name = request.form['txtnombre']
        idArea = int(request.form['txtArea'])
        cuatri = int(request.form['txtcuatri'])
        grado = int(request.form['txtGrado'])
        
        cursor = db.connection.cursor()
        cursor.execute("UPDATE grupos SET idArea=%s, nombre=%s, cuatrimestre=%s, idGrado=%s WHERE idGrupo = %s;", (idArea, name, cuatri, grado, id))
        grupo = cursor.fetchall()
        db.connection.commit()
        cursor.close()

        return redirect(url_for('grupo'))
    else:
        return protected()


@app.route('/plantillaA', methods=['POST'])
@login_required
def plantillaA():
    maestros = []
    maestros_name =[]
    datos = request.form.get
    print(datos)
    maestroOne = request.form.get('txtmaestroone', '')
    maestroTwo = request.form.get('txtmaestrotwo', '')
    maestroThree = request.form.get('txtmaestrothree', '')
    firma = request.form.get('firma', '')
    print(maestroTwo)
    maestros = []
    for maestro in [maestroOne, maestroTwo, maestroThree]:
        if maestro != '':
            maestros.append(maestro)
    
    for maestro in maestros:
        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM `maestros` WHERE idMaestro = %s;", (int(maestro),))
        maestrosN = cursor.fetchall()
        cursor.close()

        if maestrosN:  
            nombre_maestro = maestrosN[0][1]  
            apellido_maestro = maestrosN[0][2]  
            nombre_completo = f"{nombre_maestro} {apellido_maestro}"
            maestros_name.append(nombre_completo)  

    file = request.files['txtnombre']
    filename = secure_filename(file.filename)  
    timestamp = t.now().strftime('%Y%m%d%H%M%S') 
    filename_with_timestamp = timestamp + '_' + filename
    file_path = os.path.join('static/plantillas', filename_with_timestamp)
    file.save(file_path) 

    if filename.lower().endswith(('.doc', '.docx')):
        
        doc = Document(file_path)

        
        celda_actual = None

        
        for table in doc.tables:
                
                for row_idx, row in enumerate(table.rows):
                    
                    for cell_idx, cell in enumerate(row.cells):
                        
                        if 'nombre del docente' in cell.text.lower():
                            
                            if row_idx > 0:
                                celda_arriba = table.cell(row_idx - 1, cell_idx)
                                
                                if maestros_name:
                                    
                                    nombre_maestro = maestros_name.pop(0)  
                                    celda_arriba.text = nombre_maestro  
                                    
                                    for paragraph in celda_arriba.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.name = 'Arial'
                                    
                                    for paragraph in celda_arriba.paragraphs:
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        paragraph.alignment_vertical = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        
        for table in doc.tables:
                
                for row_idx, row in enumerate(table.rows):
                    
                    for cell_idx, cell in enumerate(row.cells):
                        
                        if 'nombre del director' in cell.text.lower():
                            
                            if row_idx > 0:
                                celda_arriba = table.cell(row_idx - 1, cell_idx)
                                
                                
                                celda_arriba.text = current_user.fullname  
                                
                                for paragraph in celda_arriba.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.name = 'Arial'
                                
                                for paragraph in celda_arriba.paragraphs:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    paragraph.alignment_vertical = WD_CELL_VERTICAL_ALIGNMENT.CENTER 

        
        def guardar_imagen_desde_base64(base64_string, ruta_completa_archivo):
            
            _, base64_data = base64_string.split(",", 1)

            try:
                
                image_data = base64.b64decode(base64_data)

                
                imagen = Image.open(BytesIO(image_data))

                
                imagen.save(ruta_completa_archivo, 'PNG')

                
                print("Imagen eliminada del disco.")
            except Exception as e:
                print("Error al guardar la imagen:", e)

        timestamp = t.now().strftime('%Y%m%d%H%M%S')
        nombre_archivo = timestamp+"imagen_guardada.png"
        ruta_carpeta = "static/images/"  

        
        ruta_completa_archivo = os.path.join(ruta_carpeta, nombre_archivo)

        guardar_imagen_desde_base64(firma, ruta_completa_archivo)

        
        tabla = None
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "firma del director" in cell.text.lower():
                        tabla = table
                        break
                if tabla:
                    break
            if tabla:
                break
    imagen_agregada = False
    
    if tabla:
        
        for i, row in enumerate(tabla.rows):
            for j, cell in enumerate(row.cells):
                if "firma del director" in cell.text.lower():
                    firma_cell = tabla.cell(i - 1, j)
                    firma_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  
                    firma_cell.paragraphs[0].alignment_vertical = WD_ALIGN_VERTICAL.CENTER  

                    
                    if not imagen_agregada:
                        firma_cell.paragraphs[0].add_run().add_picture(ruta_completa_archivo, width=Inches(1.0), height=Inches(0.5))
                        imagen_agregada = True  

                    break
            else:
                continue
            break
        
            print("Documento Word con firma insertada guardado:")
        else:
            print("No se encontró la tabla o el nombre 'jose uriel saenz cuellar' en el documento.")
        os.remove(ruta_completa_archivo)


        doc.save(file_path)
    else:
        
        pdf_file_path = file_path

    
        
    doc = aw.Document(file_path)
    extracted_page = doc.extract_pages(0, 1)
    extracted_page.save("static/images/" + os.path.basename(file_path) + ".jpg")

    idMateria = int(request.form['txtmateria'])
    mae = maestros[0] if len(maestros) > 0 else None
    mae2 = maestros[1] if len(maestros) > 1 else None
    mae3 = maestros[2] if len(maestros) > 2 else None

    sql = "INSERT INTO plantillas (idPlantilla, namePlantilla, idMateria, idMaestro1, idMaestro2, idMaestro3, idUser) VALUES (NULL, %s, %s, %s, %s, %s, %s);"
    datos = (os.path.basename(file_path), idMateria, mae, mae2, mae3, current_user.id,)
    cursor = db.connection.cursor()
    cursor.execute(sql, datos)
    db.connection.commit()
    cursor.close()

    return redirect(url_for('plantillas'))

@app.route('/plantilla/<int:idplantilla>', methods=['GET'])
@login_required
def delete_plantilla(idplantilla):

    cursor = db.connection.cursor()
    cursor.execute("SELECT * FROM `plantillas` WHERE `idPlantilla` = %s;", (idplantilla,))
    plantilla = cursor.fetchall()
    db.connection.commit()
    

    nombre_archivo = plantilla[0][1]
    file_path = os.path.join('static/plantillas/', nombre_archivo)
    os.remove(file_path)

    file_path = os.path.join('static/images/', nombre_archivo + '.jpg')
    os.remove(file_path)

    cursor = db.connection.cursor()
    cursor.execute("DELETE FROM `plantillas` WHERE `idPlantilla` = %s;", (idplantilla,))
    db.connection.commit()
    cursor.close()

    return redirect(url_for('plantillas'))


@app.route('/editPlantilla/<int:idPlantilla>', methods=['GET'])
@login_required
def editPlantilla(idPlantilla):
    cursor = db.connection.cursor()
    cursor.execute("""SELECT p.namePlantilla AS plantilla, p.idMaestro1 AS idM1, p.idMaestro2 AS idM2, p.idMaestro3 AS idM3, p.idMateria AS idMa, m1.Nombres AS Nombres_maestro1, m1.Apellidos AS Apellidos_maestro1,
           m2.Nombres AS Nombres_maestro2, m2.Apellidos AS Apellidos_maestro2,
           m3.Nombres AS Nombres_maestro3, m3.Apellidos AS Apellidos_maestro3,
           mt.nombre AS Nombres_materia
    FROM plantillas p
    LEFT JOIN maestros m1 ON p.idMaestro1 = m1.idMaestro
    LEFT JOIN maestros m2 ON p.idMaestro2 = m2.idMaestro
    LEFT JOIN maestros m3 ON p.idMaestro3 = m3.idMaestro
    INNER JOIN materias mt ON p.idMateria = mt.idMateria
    WHERE p.idPlantilla = %s;""", (idPlantilla,))

    plantilla_data = cursor.fetchone() 
    cursor.close()

    if not plantilla_data:
        return jsonify(), 404  

    column_names = [desc[0] for desc in cursor.description]
    plantill_dict = dict(zip(column_names, plantilla_data))

    return jsonify(plantill_dict)

@app.route('/plantillaedit', methods=['POST'])
def plantillaedit():
    maestros = []
    maestros_name = []
    maestroOne = request.form.get('txtmaestroone', '')
    maestroTwo = request.form.get('txtmaestrotwo', '')
    maestroThree = request.form.get('txtmaestrothree', '')

    for maestro in [maestroOne, maestroTwo, maestroThree]:
        if maestro != '':
            maestros.append(maestro)
 
    for maestro in maestros:
        cursor = db.connection.cursor()
        cursor.execute("SELECT * FROM `maestros` WHERE idMaestro = %s;", (int(maestro),))
        maestrosN = cursor.fetchall()
        cursor.close()

        if maestrosN:  
            nombre_maestro = maestrosN[0][1]  
            apellido_maestro = maestrosN[0][2]  
            nombre_completo = f"{nombre_maestro} {apellido_maestro}"
            maestros_name.append(nombre_completo)  

    oldfile = request.form['txtnewnombre']
    file = request.files.get('txtnombre', '') 
    
    
    filename = secure_filename(file.filename) if file else None
    timestamp = t.now().strftime('%Y%m%d%H%M%S') 
    filename_with_timestamp = timestamp + '_' + filename if filename else oldfile
    file_path = None  
    if file:
        file_path = os.path.join('static/plantillas', filename_with_timestamp)
        file.save(file_path)

        if oldfile:
            oldfile_path = os.path.join('static/plantillas', oldfile)
            if os.path.exists(oldfile_path):
                os.remove(oldfile_path)
    else:
        
        file_path = os.path.join('static/plantillas/', oldfile)

    if file and filename and filename.lower().endswith(('.doc', '.docx')):
        doc = Document(file_path)
        
        celda_actual = None
        for table in doc.tables:
            
            for row_idx, row in enumerate(table.rows):
                
                for cell_idx, cell in enumerate(row.cells):
                    
                    if 'nombre del docente' in cell.text.lower():
                        
                        if row_idx > 0:
                            celda_arriba = table.cell(row_idx - 1, cell_idx)
                            celda_arriba.text = ""  

        
        for table in doc.tables:
            
            for row_idx, row in enumerate(table.rows):
                
                for cell_idx, cell in enumerate(row.cells):
                    
                    if 'nombre del docente' in cell.text.lower():
                        
                        if row_idx > 0:
                            celda_arriba = table.cell(row_idx - 1, cell_idx)
                            
                            if maestros_name:
                                
                                nombre_maestro = maestros_name.pop(0)  
                                celda_arriba.text = nombre_maestro  

        
        doc.save(file_path)
    else:
        doc = Document(file_path)
        
        celda_actual = None

        for table in doc.tables:
            
            for row_idx, row in enumerate(table.rows):
                
                for cell_idx, cell in enumerate(row.cells):
                    
                    if 'nombre del docente' in cell.text.lower():
                        
                        if row_idx > 0:
                            celda_arriba = table.cell(row_idx - 1, cell_idx)
                            celda_arriba.text = ""  

        
        for table in doc.tables:
            
            for row_idx, row in enumerate(table.rows):
                
                for cell_idx, cell in enumerate(row.cells):
                    
                    if 'nombre del docente' in cell.text.lower():
                        
                        if row_idx > 0:
                            celda_arriba = table.cell(row_idx - 1, cell_idx)
                            
                            if maestros_name:
                                
                                nombre_maestro = maestros_name.pop(0)  
                                celda_arriba.text = nombre_maestro  

        
        doc.save(file_path)

    
    pdf_file_path = file_path

    if oldfile and oldfile != filename_with_timestamp:
        oldfile_path = os.path.join('static/plantillas', oldfile)
        newfile_path = os.path.join('static/plantillas', filename_with_timestamp)        
        oldimg_path = os.path.join('static/images', oldfile)
        os.remove(oldfile_path)
        if os.path.exists(oldfile_path):
            os.rename(oldfile_path, newfile_path)

    if file:
        doc = aw.Document(file_path)
        extracted_page = doc.extract_pages(0, 1)
        extracted_page.save("static/images/" + os.path.basename(file_path) + ".jpg")

    idMateria = int(request.form['txtmateria'])
    idPlantilla = int(request.form['id'])
    mae = maestros[0] if len(maestros) > 0 else None
    mae2 = maestros[1] if len(maestros) > 1 else None
    mae3 = maestros[2] if len(maestros) > 2 else None

    cursor = db.connection.cursor()
    cursor.execute("UPDATE plantillas SET namePlantilla=%s, idMateria=%s, idMaestro1=%s, idMaestro2=%s, idMaestro3=%s WHERE idPlantilla = %s;", (os.path.basename(file_path), idMateria, mae, mae2, mae3, idPlantilla))
    db.connection.commit()
    cursor.close()

    return redirect(url_for('plantillas'))


@app.route('/editGgrado/<int:idgrado>', methods=['GET'])
@login_required
def editGgrado(idgrado):
    cursor = db.connection.cursor()
    cursor.execute("SELECT * FROM `grado` WHERE `idGrado` = %s;", (idgrado,))
    grado_data = cursor.fetchone()  
    cursor.close()

    
    if not grado_data:
        return jsonify(), 404  

    
    
    column_names = [desc[0] for desc in cursor.description]
    grado_dict = dict(zip(column_names, grado_data))

    return jsonify(grado_dict)

@app.route('/editMaestro/<int:idmaestro>', methods=['GET'])
@login_required
def editMaestro(idmaestro):
    cursor = db.connection.cursor()
    cursor.execute("""SELECT m.*,  g.nombre, g.cuatrimestre, a.aName AS nombre_area
                   FROM `maestros` AS m
                   JOIN grupos g ON m.idGrupo = g.idGrupo
                   JOIN areas a ON g.idArea = a.idArea
                   WHERE `idMaestro` = %s;""", (idmaestro,))
    maestro_data = cursor.fetchone() 
    cursor.close()

    if not maestro_data:
        return jsonify(), 404 

    column_names = [desc[0] for desc in cursor.description]
    maestro_dict = dict(zip(column_names, maestro_data))

    return jsonify(maestro_dict)

@app.route('/editAlumnos/<int:idAlumno>', methods=['GET'])
@login_required
def editAlumno(idAlumno):
    cursor = db.connection.cursor()
    cursor.execute("""SELECT m.*,  g.nombre, g.cuatrimestre, a.aName AS nombre_area
                   FROM `students` AS m
                   JOIN grupos g ON m.idGrupo = g.idGrupo
                   JOIN areas a ON g.idArea = a.idArea
                   WHERE `idStudent` = %s;""", (idAlumno,))
    alumno_data = cursor.fetchone() 

    if not alumno_data:
        return jsonify(), 404 

    column_names = [desc[0] for desc in cursor.description]
    Alumnos_dict = dict(zip(column_names, alumno_data))
    print(Alumnos_dict)
    cursor.close()
    return jsonify(Alumnos_dict)

@app.route('/filtro', methods=['POST'])
@login_required
def filtro():
    if current_user.user_type == "Administrador":
        name = request.form['txtname']
        idArea = int(request.form['txtarea'])
        cuatri = int(request.form['txtcuatri'])
        materia = []
        planeaciones = []
        plantilla = []
        if name == 'plantillas':
            cursor = db.connection.cursor()
            cursor.execute("""SELECT plantillas.*, materias.nombre
                FROM plantillas
                INNER JOIN materias ON plantillas.idMateria = materias.idMateria
                INNER JOIN grupos AS g ON materias.idGrupo = g.idGrupo 
                INNER JOIN areas AS a ON g.idArea = a.idArea
                INNER JOIN carrera AS c ON a.idCarrera = c.idCarrera
                INNER JOIN users AS u ON c.idCarrera = u.idCarrera 
                WHERE u.idUser = %s AND g.cuatrimestre = %s AND a.idArea = %s
                ORDER BY plantillas.idPlantilla DESC;
                """, (current_user.id, cuatri, idArea,))
            result = cursor.fetchall()
        elif name == 'materias':
            cursor = db.connection.cursor()
            cursor.execute("""SELECT m.*, a.aName AS nombre_area
            FROM materias m
            JOIN grupos g ON m.idGrupo = g.idGrupo
            JOIN areas a ON g.idArea = a.idArea
            JOIN carrera c ON a.idCarrera = c.idCarrera
            JOIN users u ON c.idCarrera = u.idCarrera
            JOIN grado gr ON c.idGrado = gr.idGrado
            WHERE u.idUser = %s AND a.idArea = %s AND g.cuatrimestre = %s
            ORDER BY m.idMateria DESC;
            """, (current_user.id, cuatri, idArea,))
            materia = cursor.fetchall()
        elif name == 'planeaciones':
            cursor = db.connection.cursor()
            cursor.execute("""
                SELECT pl.*, ma.nombre , m.Nombres, m.ApellidosP
                FROM planeaciones AS pl
                INNER JOIN plantillas AS p ON p.idPlantilla = pl.idPlantilla
                INNER JOIN materias AS ma ON p.idMateria = ma.idMateria
                INNER JOIN maestros AS m ON pl.idMaestro = m.idMaestro
                INNER JOIN grupos g ON m.idGrupo = g.idGrupo
                INNER JOIN grado gr ON g.idGrado = gr.idGrado
                INNER JOIN areas a ON g.idArea = a.idArea
                INNER JOIN carrera c ON a.idCarrera = c.idCarrera
                WHERE p.idUser = %s AND a.idArea = %s AND g.cuatrimestre = %s
                ORDER BY pl.idPlan DESC;
            """, (current_user.id, cuatri, idArea,))
            planeaciones = cursor.fetchall()
        elif name == 'plantillas':
            cursor = db.connection.cursor()
            cursor.execute("""SELECT plantillas.*, materias.nombre
                FROM plantillas
                INNER JOIN materias ON plantillas.idMateria = materias.idMateria
                INNER JOIN grupos AS g ON materias.idGrupo = g.idGrupo 
                INNER JOIN areas AS a ON g.idArea = a.idArea
                INNER JOIN carrera AS c ON a.idCarrera = c.idCarrera
                INNER JOIN users AS u ON c.idCarrera = u.idCarrera 
                WHERE u.idUser = %s AND g.cuatrimestre = %s AND a.idArea = %s
                ORDER BY plantillas.idPlantilla DESC;
                """, (current_user.id, cuatri, idArea,))
        
        db.connection.commit()
        cursor.close()
        
        return render_template('admin/home.html', plantilla=result, planeaciones=planeaciones, materia=materia)
    else:
        return protected()


@app.route('/editpl', methods=['POST'])
@login_required
def editpl():
    name = request.form['txtfile']
    doc_path = os.path.join('static/planeacion/', name)

    timestamp = t.now().strftime('%Y%m%d%H%M%S') 
    filename_with_timestamp = timestamp + '_' + name
    carpeta_modificados = os.path.join('static/planeacion/', filename_with_timestamp)
    cuadros_encontrados = 0
    resultados = {
        "cuadros_porcentaje": [],
        "cuadros_actividades": [],
        "cuadros_fechas": [],
        "cuadros_observaciones": [],
        "cuadros_temas": [],
        "cuadros_conceptos": [],
        "conocimientos_generales": [],
        "cuadro_subtemas": []
    }

    doc = docx.Document(doc_path)

    for i, table in enumerate(doc.tables, 1):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                if "fecha planeada por semana".lower() in cell.text.lower():
                    cuadros_encontrados += 1
                    texto_cuadros_debajo = []
                    for j in range(1, 5):
                        if row_index + j < len(table.rows):
                            siguiente_fila = table.rows[row_index + j]
                            texto_cuadro_debajo = siguiente_fila.cells[cell_index].text
                            texto_cuadros_debajo.append(texto_cuadro_debajo)
                    resultados["cuadros_fechas"].append((cuadros_encontrados, texto_cuadros_debajo))

                if "observaciones a la materia" in cell.text.lower():
                    cuadros_encontrados += 1
                    texto_cuadros_debajo = []
                    for j in range(1, 2):
                        if row_index + j < len(table.rows):
                            siguiente_fila = table.rows[row_index + j]
                            texto_cuadro_debajo = siguiente_fila.cells[cell_index].text
                            texto_cuadros_debajo.append(texto_cuadro_debajo)
                    resultados["cuadros_observaciones"].append((i, texto_cuadros_debajo))

                if "temas" in cell.text.lower():
                    cuadros_encontrados += 1
                    texto_cuadros_debajo = []
                    for j in range(1, 4):
                        if row_index + j < len(table.rows):
                            siguiente_fila = table.rows[row_index + j]
                            texto_cuadro_debajo = siguiente_fila.cells[cell_index].text
                            texto_cuadros_debajo.append(texto_cuadro_debajo)
                    resultados["cuadros_temas"].append((i, texto_cuadros_debajo))
                
                if "subtema" in cell.text.lower():
                    cuadros_encontrados += 1
                    texto_cuadros_debajo = []
                    for j in range(1, 4):
                        if row_index + j < len(table.rows):
                            siguiente_fila = table.rows[row_index + j]
                            texto_cuadro_debajo = siguiente_fila.cells[cell_index].text
                            texto_cuadros_debajo.append(texto_cuadro_debajo)
                    resultados["cuadro_subtemas"].append((i, texto_cuadros_debajo))

                if "concepto" in cell.text.lower():
                    cuadros_encontrados += 1
                    texto_cuadros_debajo = []
                    for j in range(1, 5):
                        if row_index + j < len(table.rows):
                            siguiente_fila = table.rows[row_index + j]
                            texto_cuadro_debajo = siguiente_fila.cells[cell_index].text
                            texto_cuadros_debajo.append(texto_cuadro_debajo)
                    resultados["cuadros_conceptos"].append((i, texto_cuadros_debajo))
                
                if "conocimientos generales del profesor" in cell.text.lower():
                    cuadros_encontrados += 1
                    texto_cuadros_debajo = []
                    for j in range(1, 2):
                        if row_index + j < len(table.rows):
                            siguiente_fila = table.rows[row_index + j]
                            texto_cuadro_debajo = siguiente_fila.cells[cell_index].text
                            texto_cuadros_debajo.append(texto_cuadro_debajo)
                    resultados["conocimientos_generales"].append((i, texto_cuadros_debajo))

    def obtener_datos_cuarta_columna_despues_de_concepto(doc_path):
        datos_tercera_columna = []

        palabra_clave_concepto = "concepto"

        ultima_tabla = None
        for table in doc.tables:
            ultima_tabla = table

        if ultima_tabla:
            
            posicion_concepto = None
            for fila_index, fila in enumerate(ultima_tabla.rows):
                celda = fila.cells[3]  
                if palabra_clave_concepto.lower() in celda.text.lower():
                    posicion_concepto = (fila_index, 3)  
                    break  
            
            if posicion_concepto:
                fila_concepto, columna_concepto = posicion_concepto
                numero_filas = len(ultima_tabla.rows)

                for fila in range(fila_concepto + 1, numero_filas):
                    celda = ultima_tabla.rows[fila].cells[3]  
                    datos_celda = celda.text.strip()
                    datos_tercera_columna.append((fila, datos_celda))  

        return datos_tercera_columna


    resulta = obtener_datos_cuarta_columna_despues_de_concepto(doc_path)

    def obtener_datos_cuarta_columna_despues_de_subtems(doc_path):
        datos_tercera_columna = []

        palabra_clave_concepto = "concepto"

        ultima_tabla = None
        for table in doc.tables:
            ultima_tabla = table

        if ultima_tabla:
            
            posicion_concepto = None
            for fila_index, fila in enumerate(ultima_tabla.rows):
                celda = fila.cells[3]  
                if palabra_clave_concepto.lower() in celda.text.lower():
                    posicion_concepto = (fila_index, 3)  
                    break  
            
            if posicion_concepto:
                fila_concepto, columna_concepto = posicion_concepto
                numero_filas = len(ultima_tabla.rows)

                for fila in range(fila_concepto + 1, numero_filas):
                    celda = ultima_tabla.rows[fila].cells[2]  
                    datos_celda = celda.text.strip()
                    datos_tercera_columna.append((fila, datos_celda))  

        return datos_tercera_columna


    resultaSub = obtener_datos_cuarta_columna_despues_de_subtems(doc_path)

    def obtener_datos_cuarta_columna_despues_de_subtems(doc_path):
        datos_tercera_columna = []

        palabra_clave_concepto = "concepto"

        ultima_tabla = None
        for table in doc.tables:
            ultima_tabla = table

        if ultima_tabla:
            
            posicion_concepto = None
            for fila_index, fila in enumerate(ultima_tabla.rows):
                celda = fila.cells[3]  
                if palabra_clave_concepto.lower() in celda.text.lower():
                    posicion_concepto = (fila_index, 3)  
                    break  
            
            if posicion_concepto:
                fila_concepto, columna_concepto = posicion_concepto
                numero_filas = len(ultima_tabla.rows)

                for fila in range(fila_concepto + 1, numero_filas):
                    celda = ultima_tabla.rows[fila].cells[1]  
                    datos_celda = celda.text.strip()
                    datos_tercera_columna.append((fila, datos_celda))  

        return datos_tercera_columna


    resultaCon = obtener_datos_cuarta_columna_despues_de_subtems(doc_path)

    combinaciones_unicas = {}
    for num, tema in resultaCon:
        if tema not in combinaciones_unicas or num < combinaciones_unicas[tema][0]:
            combinaciones_unicas[tema] = (num, tema)

    
    resultadoCON = list(combinaciones_unicas.values())
    print(resultadoCON)
    
    def extraer_datos_de_tabla(doc_path):
        datos_tabla = []

        
        if len(doc.tables) > 0:
            ultima_tabla = doc.tables[-1]

            
            for fila_index, fila in enumerate(ultima_tabla.rows[2:], start=2):
                celdas = fila.cells

                
                datos_fila = (celdas[1].text, celdas[2].text)
                datos_tabla.append(datos_fila)

        return datos_tabla

    def agrupar_datos_por_tema(datos):
        temas_agrupados = defaultdict(list)

        for tupla in datos:
            tema = tupla[0]
            descripcion = tupla[1]
            temas_agrupados[tema].append(descripcion)

        return temas_agrupados

    
    datos_extraidos = extraer_datos_de_tabla(doc_path)

    
    datos_agrupados = agrupar_datos_por_tema(datos_extraidos)

    
    resultados_data = []
    for tema, descripciones in datos_agrupados.items():
        resultados_data.append((tema, '\n'.join(descripciones)))  
    print(resultados_data)

    def buscar_temas_en_tabla(doc_path):
        cuadros_encontrados = 0
        resultados = []

        doc = docx.Document(doc_path)

        unidad_tematica = None  

        for i, table in enumerate(doc.tables, 1):
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    if cell_index + 1 < len(row.cells):
                        celda_derecha = row.cells[cell_index + 1]
                        if "unidad temática" in cell.text.lower():
                            unidad_tematica = celda_derecha
                            break
                
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    if "temas" in cell.text.lower():
                        cuadros_encontrados += 1

                        if unidad_tematica:
                            texto_cuadros_debajo = []
                            for j in range(1, 4):
                                if row_index + j < len(table.rows):
                                    siguiente_fila = table.rows[row_index + j]
                                    texto_cuadro_debajo = siguiente_fila.cells[cell_index].text
                                    texto_cuadros_debajo.append(texto_cuadro_debajo)
                            resultados.append((unidad_tematica.text, i, texto_cuadros_debajo))

        return cuadros_encontrados, resultados

    cuadros_encontrados, resultr = buscar_temas_en_tabla(doc_path)
    print(resultr)
    
    def obtener_datos_tabla(doc_path):
        doc = docx.Document(doc_path)
        datos = []

        for table in doc.tables:
            descripcion_columna = None
            porcentaje_columna = None

            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    texto = cell.text.lower()
                    
                    if "descripción de la actividad para alcanzar el resultado de aprendizaje de la unidad" in texto:
                        descripcion_columna = cell_index
                    elif "% de la evaluación de la unidad" in texto:
                        porcentaje_columna = cell_index

                if descripcion_columna is not None and porcentaje_columna is not None:
                    break

            if descripcion_columna is not None and porcentaje_columna is not None:
                for row in table.rows[row_index + 1:]:  
                    descripcion = row.cells[descripcion_columna].text
                    porcentaje = row.cells[porcentaje_columna].text
                    if descripcion.strip() and porcentaje.strip():  
                        datos.append((descripcion, porcentaje))

                break  

        return datos

    datos_tabla = obtener_datos_tabla(doc_path)
    
    formatted_dataA = [item for _, items_list in resultados["cuadros_temas"] for item in items_list]
    formatted_dataAC = [item for _, items_list in resultados["cuadros_actividades"] for item in items_list]
    formatted_dataPR = [item for _, items_list in resultados["cuadros_porcentaje"] for item in items_list]
    formatted_dataCON = [item for _, items_list in resultados["cuadros_conceptos"] for item in items_list]
    formatted_dataSUB = [item for _, items_list in resultados["cuadros_temas"] for item in items_list]
    formatted_dataOROLD = [item for _, items_list in resultados["cuadros_observaciones"] for item in items_list]
    formatted_dataOR = [texto.replace('\n', '') for texto in formatted_dataOROLD]
    formatted_dataGENOLD = [item for _, items_list in resultados["conocimientos_generales"] for item in items_list]
    formatted_dataGEN = [texto.replace('\n', '') for texto in formatted_dataGENOLD]
    resultado = [(formatted_dataSUB[numero[0] - 2], numero[1]) for numero in resulta]


    def extract_dates_and_weeks(strings):
        data_list = []
        for s in strings:
            
            parts = s.split('\n')
            
            if len(parts) >= 4 and parts[1] and parts[3]:
                
                week = int(parts[0].split()[1])
                start_date = t.strptime(parts[1], '%Y-%m-%d')
                end_date = t.strptime(parts[3], '%Y-%m-%d')
                
                data_list.append((week, start_date, end_date))
            else:
                print("Error: Datos faltantes o formato incorrecto en la cadena:", s)
        
        return data_list

    all_data = []
    if resultados["cuadros_fechas"]:
        all_data = []
        for number, strings in resultados["cuadros_fechas"]:
            extracted_data = extract_dates_and_weeks(strings)
            all_data.extend(extracted_data)

        
        print(all_data)
    else:
        print("No hay datos en la lista 'resultados[cuadros_fechas]'.")

    formatted_dataF = [(week, start_date.strftime('%Y-%m-%d'), end_date.strftime('%Y-%m-%d')) for week, start_date, end_date in all_data]
    
    resultados = [(fecha[0], tema, fecha[1], fecha[2]) for fecha, tema in zip(formatted_dataF, resultr)]
    print(formatted_dataAC)
    
    
    return render_template('user/formEdit.html', name=name, all_data=resultados, formatted_dataAC=datos_tabla, formatted_dataPR=formatted_dataPR, formatted_dataOR=formatted_dataOR, formatted_dataGEN=formatted_dataGEN, resultado=resultado, datos_agrupados=resultados_data)


@app.route('/filtro_admin', methods=['POST'])
@login_required
def filtro_admin():
    if current_user.user_type == "Superusuario":
        name = request.form['txtname']
        grado = []
        edificio = []
        special = []
        admin = []
        career = []
        spe = []
        if name == "carrera":
            cursor = db.connection.cursor()
            cursor.execute("SELECT carrera.*, grado.nameGrado, edificio.nameE FROM carrera INNER JOIN grado ON carrera.idGrado = grado.idGrado INNER JOIN edificio ON carrera.idEdificio = edificio.idEdificio ORDER BY idCarrera DESC;")
            career = cursor.fetchall()
            cursor = db.connection.cursor()
            cursor.execute("SELECT areas.*, carrera.carreraName FROM areas INNER JOIN carrera ON areas.idCarrera = carrera.idCarrera ORDER BY idArea DESC;")
            spe = cursor.fetchall()
        elif name == "users":
            cursor = db.connection.cursor()
            cursor.execute("SELECT * FROM users ORDER BY idUser DESC;")
            admin = cursor.fetchall()
        elif name == "areas":
            cursor = db.connection.cursor()
            cursor.execute("SELECT areas.*, carrera.carreraName FROM areas INNER JOIN carrera ON areas.idCarrera = carrera.idCarrera ORDER BY idArea DESC;")
            special = cursor.fetchall()
        elif name == "edificio":
            cursor = db.connection.cursor()
            cursor.execute("SELECT * FROM edificio ORDER BY idEdificio DESC;")
            edificio = cursor.fetchall()
        elif name == "grado":
            cursor = db.connection.cursor()
            cursor.execute("SELECT * FROM grado ORDER BY idGrado DESC;")
            grado = cursor.fetchall()

        cursor.close()
        
        return render_template('superAdmin/home.html', carrera=career, admin=admin, special=special, edificio=edificio, grado=grado, spe=spe)
    else:
        return protected()


@app.route('/delete_plan', methods=['POST'])
@login_required
def delete_plan():
    id = request.form['id']
    cursor = db.connection.cursor()
    cursor.execute("SELECT name FROM planeaciones WHERE idPlan = %s;", (id,))
    current_token = cursor.fetchone()
    file_path = os.path.join('static/planeacion/', current_token[0])
    os.remove(file_path)
    cursor = db.connection.cursor()
    cursor.execute("DELETE FROM `planeaciones` WHERE `idPlan` = %s;", (id,))
    db.connection.commit()
    cursor.close()


    return redirect(url_for('plan'))

@app.route('/alumnoDelete/<int:idAlumno>', methods=['GET'])
@login_required
def alumnoDelete(idAlumno):
    cursor = db.connection.cursor()
    cursor.execute("DELETE FROM `students` WHERE `idStudent` = %s;", (idAlumno,))
    db.connection.commit()
    cursor.close()


    return redirect(url_for('alumnos'))

@app.route('/lock', methods=['POST'])
@login_required
def lock():
    id = request.form['id']
    print(id)
    cursor = db.connection.cursor()
    cursor.execute("SELECT token FROM planeaciones WHERE idPlan = %s;", (id,))
    current_token = cursor.fetchone()
    print(current_token)
    
    if current_token[0] == "ACCESS":
        new_token = "UNACCESS"  
    else:
        new_token = "ACCESS"

    cursor = db.connection.cursor()
    cursor.execute("UPDATE planeaciones SET token = %s WHERE idPlan = %s;", (new_token, id,))
    db.connection.commit()
    cursor.close()


    return redirect(url_for('plan'))


@app.route('/planeacionEdit', methods=['POST'])
@login_required
def planeacionEdit():
    if current_user.user_type == "Maestro":
        
        data = request.form
        
        name = data.get('name')

        def buscar_porcentajes_y_reemplazar(doc_path):
            timestamp = t.now().strftime('%Y%m%d%H%M%S') 
            filename_with_timestamp = timestamp + '_' + name
            carpeta_modificados = os.path.join('static/planeacion/', filename_with_timestamp)
            cuadros_encontrados = 0
            resultados = []
            porcentaje = [data.get(f'porcentaje_{i}', '') for i in range(1, 5)]
            actividades = [data.get(f'actividad_{i}', '') for i in range(1, 5)]
            semana = [data.get(f'numero_semana_{i}', '') for i in range(1, 15)]
            fecha_inicio = [data.get(f'fecha_inicio_{i}', '') for i in range(1, 15)]
            fecha_final = [data.get(f'fecha_fin_{i}', '') for i in range(1, 15)]
            conceptos = [data.get(f'conceptos_t_{i}', '') for i in range(1, 15)]
            temas_a_usar = [data.get(f'temas_{i}', '') for i in range(1, 15)]
            firma_base64 = data.get('firmas', '')  
            observacion = data.get('observacion', '')
            conocimientos = data.get('conocimientos', '')
            print(firma_base64)
            
            doc = docx.Document(doc_path)

            
            for i, table in enumerate(doc.tables, 1):
                for row_index, row in enumerate(table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        
                        if "% de la evaluación de la unidad" in cell.text.lower():
                            cuadros_encontrados += 1
                            
                            texto_cuadros_debajo = []  
                            for j in range(1, 5):  
                                if row_index + j < len(table.rows):
                                    siguiente_fila = table.rows[row_index + j]
                                    indice_texto = (j - 1) % len(porcentaje)  
                                    texto_cuadro_debajo = siguiente_fila.cells[cell_index].text
                                    texto_cuadro_debajo = texto_cuadro_debajo if texto_cuadro_debajo is not None else ""
                                    if porcentaje[indice_texto] != '':
                                        siguiente_fila.cells[cell_index].text = porcentaje[indice_texto]
                                    else:
                                        siguiente_fila.cells[cell_index].text = ""
                                    texto_cuadros_debajo.append(texto_cuadro_debajo)
                                    
                                    for paragraph in siguiente_fila.cells[cell_index].paragraphs:
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        for run in paragraph.runs:
                                            run.font.name = 'Arial'
                                    
                                    siguiente_fila.cells[cell_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                else:
                                    texto_cuadros_debajo.append("")  
                            print(texto_cuadros_debajo)
                            resultados.append((i, texto_cuadros_debajo))

            
            for i, table in enumerate(doc.tables, 1):
                for row_index, row in enumerate(table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        
                        if "descripción de la actividad para alcanzar el resultado de aprendizaje de la unidad" in cell.text.lower():
                            cuadros_encontrados += 1
                            
                            texto_cuadros_debajo = []  
                            for j in range(1, 5):  
                                if row_index + j < len(table.rows):
                                    siguiente_fila = table.rows[row_index + j]
                                    indice_texto = (j - 1) % len(actividades)  
                                    texto_cuadro_debajo = siguiente_fila.cells[cell_index].text
                                    if actividades[indice_texto] != '':
                                        siguiente_fila.cells[cell_index].text = actividades[indice_texto]
                                        for paragraph in siguiente_fila.cells[cell_index].paragraphs:
                                            for run in paragraph.runs:
                                                run.font.name = 'Arial'
                                    else:
                                        siguiente_fila.cells[cell_index].text = ""
                                    texto_cuadros_debajo.append(texto_cuadro_debajo)
                                    
                                    
                                    siguiente_fila.cells[cell_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                else:
                                    texto_cuadros_debajo.append("")
                            
                            resultados.append((i, texto_cuadros_debajo))

            cuadros_encontrados_totales = 0

            for i, table in enumerate(doc.tables, 1):
                for row_index, row in enumerate(table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        
                        if "fecha planeada por semana".lower() in cell.text.lower():
                            cuadros_encontrados_totales += 1
                            
                            texto_cuadros_debajo = []
                            for j in range(1, 5):  
                                if row_index + j < len(table.rows):
                                    siguiente_fila = table.rows[row_index + j]
                                    texto_cuadro_debajo = siguiente_fila.cells[cell_index].text
                                    texto_cuadros_debajo.append(texto_cuadro_debajo)
                                    siguiente_fila.cells[cell_index].text = ""  

                            resultados.append((cuadros_encontrados_totales, texto_cuadros_debajo))

                            
                            for j, texto_reemplazo in enumerate(semana):
                                if j < len(texto_cuadros_debajo):  
                                    siguiente_fila = table.rows[row_index + j + 1]
                                    siguiente_celda = siguiente_fila.cells[cell_index]
                                    if texto_reemplazo != '' and fecha_inicio[j] != '' and fecha_final[j] != '' and semana[j] != '':
                                        siguiente_celda.text = f"Semana {semana[j]}\n{fecha_inicio[j]}\nA\n{fecha_final[j]}"
                                        parrafo = siguiente_celda.paragraphs[0]
                                        run = parrafo.runs[0]
                                        font = run.font
                                        font.name = 'Arial'
                                        parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                        siguiente_celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                    else:
                                        siguiente_celda.text = ""

                            semana = semana[len(texto_cuadros_debajo):]  
                            fecha_inicio = fecha_inicio[len(texto_cuadros_debajo):]  
                            fecha_final = fecha_final[len(texto_cuadros_debajo):]  

            for i, table in enumerate(doc.tables, 1):
                for row_index, row in enumerate(table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        
                        if "observaciones a la materia" in cell.text.lower():
                            cuadros_encontrados += 1
                            
                            texto_cuadros_debajo = []
                            for j in range(1, 2):  
                                if row_index + j < len(table.rows):
                                    siguiente_fila = table.rows[row_index + j]
                                    siguiente_fila.cells[cell_index].text = observacion + "\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n"
                                    for paragraph in siguiente_fila.cells[cell_index].paragraphs:
                                        for run in paragraph.runs:
                                            run.font.name = 'Arial'
                                    
                            
                            resultados.append((i, texto_cuadros_debajo))

            def buscar_temas_en_tabla(doc_path):
                cuadros_encontrados = 0
                resultados = []

                doc = docx.Document(doc_path)

                for i, table in enumerate(doc.tables, 1):
                    for row_index, row in enumerate(table.rows):
                        for cell_index, cell in enumerate(row.cells):
                            
                            if "temas" in cell.text.lower():
                                cuadros_encontrados += 1
                                
                                texto_cuadros_debajo = []
                                for j in range(1, 4):  
                                    if row_index + j < len(table.rows):
                                        siguiente_fila = table.rows[row_index + j]
                                        texto_cuadro_debajo = siguiente_fila.cells[cell_index].text
                                        texto_cuadros_debajo.append(texto_cuadro_debajo)
                                resultados.append((i, texto_cuadros_debajo))

                

                return cuadros_encontrados, resultados

            cuadros_encontrados, resultados = buscar_temas_en_tabla(doc_path)

            contador = 2
            
            resultado_formateado = []
            for inicio, temas in resultados:
                for tema in temas:
                    resultado_formateado.append((contador, tema))
                    contador += 1


            def obtener_datos_tercera_columna_despues_de_subtema(doc_path):
                datos_tercera_columna = []

                palabra_clave_subtema = "subtema"

                ultima_tabla = None
                for table in doc.tables:
                    ultima_tabla = table

                if ultima_tabla:
                    
                    posicion_subtema = None
                    for fila_index, fila in enumerate(ultima_tabla.rows):
                        for celda_index, celda in enumerate(fila.cells):
                            if palabra_clave_subtema.lower() in celda.text.lower():
                                posicion_subtema = (fila_index, celda_index)

                    
                    if posicion_subtema:
                        fila_subtema, columna_subtema = posicion_subtema
                        numero_filas = len(ultima_tabla.rows)

                        for fila in range(fila_subtema + 1, numero_filas):
                            celda = ultima_tabla.rows[fila].cells[2]  
                            datos_celda = celda.text.strip()
                            datos_tercera_columna.append((fila, datos_celda))  

                return datos_tercera_columna, ultima_tabla

            
            datos_tercera_columna_despues_subtemas, ultima_tabla = obtener_datos_tercera_columna_despues_de_subtema(doc_path)

            datos_tercera_columna_despues_subtema = resultado_formateado
            
            for fila, dato in datos_tercera_columna_despues_subtema:
                nuevo_dato = dato 
                celda_editar = ultima_tabla.rows[fila].cells[2]
                celda_editar.text = nuevo_dato
                for paragraph in ultima_tabla.rows[fila].cells[2].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)

            def obtener_datos_cuarta_columna_despues_de_concepto(doc_path):
                datos_tercera_columna = []

                palabra_clave_concepto = "concepto"

                ultima_tabla = None
                for table in doc.tables:
                    ultima_tabla = table

                if ultima_tabla:
                    
                    posicion_concepto = None
                    for fila_index, fila in enumerate(ultima_tabla.rows):
                        celda = fila.cells[3]  
                        if palabra_clave_concepto.lower() in celda.text.lower():
                            posicion_concepto = (fila_index, 3)  
                            break  

                    
                    if posicion_concepto:
                        fila_concepto, columna_concepto = posicion_concepto
                        numero_filas = len(ultima_tabla.rows)

                        for fila in range(fila_concepto + 1, numero_filas):
                            celda = ultima_tabla.rows[fila].cells[2]  
                            datos_celda = celda.text.strip()
                            datos_tercera_columna.append((fila, datos_celda))  

                return datos_tercera_columna, ultima_tabla


            resulta, ultima_tabla = obtener_datos_cuarta_columna_despues_de_concepto(doc_path)

            

            datos_contados = len(resulta)

            
            datos_formateados = []
            indice_deseado = 2

            for indice, dato in enumerate(conceptos):
                datos_formateados.append((indice_deseado, dato))
                indice_deseado += 1

                
                if indice_deseado > datos_contados + 1:  
                    break


            for fila, dato in datos_formateados:
                nuevo_dato = dato 
                celda_editar = ultima_tabla.rows[fila].cells[3]  
                celda_editar.text = nuevo_dato
                for paragraph in ultima_tabla.rows[fila].cells[3].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)

            def obtener_celda_despues_de_conocimientos_generales(doc_path):
                palabra_clave_concepto = "conocimientos generales del profesor"
                celda_despues_conocimientos = None

                
                ultima_tabla = doc.tables[-1]  

                for fila_index, fila in enumerate(ultima_tabla.rows):
                    for celda_index, celda in enumerate(fila.cells):
                        if palabra_clave_concepto.lower() in celda.text.lower():
                            
                            if fila_index + 1 < len(ultima_tabla.rows):
                                celda_despues_conocimientos = ultima_tabla.rows[fila_index + 1].cells[0]  
                            break
                    if celda_despues_conocimientos is not None:
                        break

                return celda_despues_conocimientos, ultima_tabla

            
            celda_despues_conocimientos, ultima_tabla = obtener_celda_despues_de_conocimientos_generales(doc_path)

            
            if celda_despues_conocimientos is not None:
                nuevo_dato = conocimientos  
                celda_despues_conocimientos.text = nuevo_dato
                primer_run = celda_despues_conocimientos.paragraphs[0].runs[0]
                primer_run.font.name = 'Arial'

            
            
            def obtener_dato_columna(doc_path):

                datos_tercera_columna = []

                palabra_clave_subtema = "tema"

                ultima_tabla = None
                for table in doc.tables:
                    ultima_tabla = table

                if ultima_tabla:
                    posicion_subtema = None
                    for fila_index, fila in enumerate(ultima_tabla.rows):
                        for celda_index, celda in enumerate(fila.cells):
                            if palabra_clave_subtema.lower() in celda.text.lower():
                                posicion_subtema = (fila_index, celda_index)

                    if posicion_subtema:
                        fila_subtema, columna_subtema = posicion_subtema
                        numero_filas = len(ultima_tabla.rows)

                        for fila in range(fila_subtema + 1, numero_filas):
                            celda = ultima_tabla.rows[fila].cells[1]
                            datos_celda = celda.text.strip()
                            datos_tercera_columna.append((fila, datos_celda))

                return datos_tercera_columna, ultima_tabla

            datos_tercera_columna_despues_subtemas, ultima_tabla = obtener_dato_columna(doc_path)

            datos_dict = {}

            
            for tupla in datos_tercera_columna_despues_subtemas:
                datos_dict[tupla[1]] = tupla

            
            datos_tercera_columna_despues_subtemas_sin_repetidos = list(datos_dict.values())

            
            temas = temas_a_usar
            datos_contados = len(datos_tercera_columna_despues_subtemas_sin_repetidos)

            datos_formateados = []

            
            for fila, (indice, dato) in enumerate(datos_tercera_columna_despues_subtemas_sin_repetidos, start=2):
                
                if fila - 2 < len(temas):
                    nuevo_dato = temas[fila - 2]
                else:
                    nuevo_dato = ""
                
                
                datos_formateados.append((indice, nuevo_dato))

            print(datos_formateados)
            
            for fila, dato in datos_formateados:
                nuevo_dato = dato
                celda_editar = ultima_tabla.rows[fila].cells[1]
                celda_editar.text = nuevo_dato
                for paragraph in ultima_tabla.rows[fila].cells[1].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
                        run.bold = True


                def guardar_imagen_desde_base64(base64_string, ruta_completa_archivo):
                    
                    _, base64_data = base64_string.split(",", 1)

                    try:
                        
                        image_data = base64.b64decode(base64_data)

                        
                        imagen = Image.open(BytesIO(image_data))

                        
                        imagen.save(ruta_completa_archivo, 'PNG')

                        
                        print("Imagen eliminada del disco.")
                    except Exception as e:
                        print("Error al guardar la imagen:", e)

                nombre_archivo = "imagen_guardada.png"
                ruta_carpeta = "static/images/"  

                
                ruta_completa_archivo = os.path.join(ruta_carpeta, nombre_archivo)

                guardar_imagen_desde_base64(firma_base64, ruta_completa_archivo)

                
                tabla = None
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if current_user.fullname.lower() in cell.text.lower():
                                tabla = table
                                break
                        if tabla:
                            break
                    if tabla:
                        break
            imagen_agregada = False
            
            if tabla:
                
                for i, row in enumerate(tabla.rows):
                    for j, cell in enumerate(row.cells):
                        if current_user.fullname.lower() in cell.text.lower():
                            firma_cell = tabla.cell(i - 2, j)
                            
                            
                            for paragraph in firma_cell.paragraphs:
                                for run in paragraph.runs:
                                    run.clear()
                            
                            firma_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  
                            firma_cell.paragraphs[0].alignment_vertical = WD_ALIGN_VERTICAL.CENTER  

                            if not imagen_agregada:
                                firma_cell.paragraphs[0].add_run().add_picture(ruta_completa_archivo, width=Inches(1.0), height=Inches(0.5))
                                imagen_agregada = True  

                            break
                    else:
                        continue
                    break

                
                    print("Documento Word con firma insertada guardado:")
                else:
                    print("No se encontró la tabla o el nombre 'jose uriel saenz cuellar' en el documento.")
                os.remove(ruta_completa_archivo)

            ruta_archivo_modificado = os.path.join('static/planeacion/', name)

            with open(ruta_archivo_modificado, "wb") as f:
                doc.save(f)

        doc_path = os.path.join('static/planeacion/', name)

        buscar_porcentajes_y_reemplazar(doc_path)
        acc = 'UNACCESS'
        current_datetime = datetime.today().strftime('%Y-%m-%d')
        cursor = db.connection.cursor()
        cursor.execute("UPDATE planeaciones SET fechaedit = %s, token=%s WHERE name = %s;", (current_datetime, acc, name,))
        db.connection.commit()
        cursor.close()
        return redirect(url_for('planeacionesU'))  
    

    else:
        return protected()


def status_401(error):
    return redirect(url_for('login'))

def status_404(error):
    return "<h1>Página no encontrada</h1>", 404



if __name__ == '__main__':
    app.config.from_object(config['DevelopmentConfig'])
    app.register_error_handler(401, status_401)
    app.register_error_handler(404, status_404)
    app.run(host="0.0.0.0", port=5000, debug=False)






